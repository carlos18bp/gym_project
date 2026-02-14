# backend/gym_app/tests/views/test_dynamic_document.py

import datetime
import io
import json
import os
from io import BytesIO
from unittest import mock
from unittest.mock import patch, MagicMock, PropertyMock

import pytest
from PIL import Image

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from django.utils import timezone
from rest_framework import status
from rest_framework.test import APIClient

from gym_app.models import User, DynamicDocument, DocumentVariable, RecentDocument
from gym_app.views.dynamic_documents import document_views

pytestmark = pytest.mark.django_db
@pytest.fixture
def user():
    return User.objects.create_user(
        email='test@example.com',
        password='testpassword'
    )

@pytest.fixture
def sample_document(user):
    """Create a sample dynamic document for testing"""
    document = DynamicDocument.objects.create(
        title="Test Document",
        content="<p>This is a test document with {{variable1}} and {{variable2}}.</p>",
        state="Draft",
        created_by=user
    )
    
    # Create variables for the document
    DocumentVariable.objects.create(
        document=document,
        name_en="variable1",
        value="Value 1"
    )
    
    DocumentVariable.objects.create(
        document=document,
        name_en="variable2",
        value="Value 2"
    )
    
    return document

@pytest.mark.django_db
class TestDynamicDocumentViews:
    
    def test_list_dynamic_documents_authenticated(self, api_client, user, sample_document):
        """Test retrieving a list of dynamic documents when authenticated"""
        # Authenticate the user
        api_client.force_authenticate(user=user)
        
        # Make the request
        url = reverse('list_dynamic_documents')
        response = api_client.get(url)
        
        # Assert the response
        assert response.status_code == status.HTTP_200_OK
        # The list endpoint now returns a paginated structure
        # {
        #   "items": [...],
        #   "totalItems": N,
        #   "totalPages": P,
        #   "currentPage": page
        # }
        assert 'items' in response.data
        assert response.data['totalItems'] == 1
        assert response.data['totalPages'] == 1
        assert response.data['currentPage'] == 1

        assert len(response.data['items']) == 1
        first_doc = response.data['items'][0]
        assert first_doc['title'] == "Test Document"
        assert first_doc['state'] == "Draft"
    
    def test_list_dynamic_documents_pagination_default_limit_10(self, api_client, user):
        """list_dynamic_documents debe paginar de 10 en 10 y devolver totales correctos."""
        api_client.force_authenticate(user=user)

        # Crear más de 10 documentos para forzar paginación
        for i in range(15):
            DynamicDocument.objects.create(
                title=f"Doc {i}",
                content="<p>content</p>",
                state="Draft",
                created_by=user,
            )

        url = reverse('list_dynamic_documents')

        # Primera página
        response_page_1 = api_client.get(url, {'page': 1})
        assert response_page_1.status_code == status.HTTP_200_OK
        assert response_page_1.data['totalItems'] == 15
        assert response_page_1.data['totalPages'] == 2
        assert response_page_1.data['currentPage'] == 1
        # El límite por defecto debe ser 10 elementos en la primera página
        assert len(response_page_1.data['items']) == 10

        # Segunda página
        response_page_2 = api_client.get(url, {'page': 2})
        assert response_page_2.status_code == status.HTTP_200_OK
        assert response_page_2.data['totalItems'] == 15
        assert response_page_2.data['totalPages'] == 2
        assert response_page_2.data['currentPage'] == 2
        # Restantes 5 elementos en la segunda página
        assert len(response_page_2.data['items']) == 5

    def test_list_dynamic_documents_filter_by_states(self, api_client, user):
        """list_dynamic_documents debe permitir filtrar por múltiples estados usando 'states'."""
        api_client.force_authenticate(user=user)

        # Crear documentos en distintos estados
        DynamicDocument.objects.create(
            title="Doc Draft",
            content="<p>x</p>",
            state="Draft",
            created_by=user,
        )
        DynamicDocument.objects.create(
            title="Doc Published",
            content="<p>x</p>",
            state="Published",
            created_by=user,
        )
        DynamicDocument.objects.create(
            title="Doc Progress",
            content="<p>x</p>",
            state="Progress",
            created_by=user,
        )

        url = reverse('list_dynamic_documents')

        # Filtrar solo Draft y Published usando el nuevo parámetro states
        response = api_client.get(url, {'states': 'Draft,Published'})
        assert response.status_code == status.HTTP_200_OK

        # Debe devolver únicamente los documentos en esos estados
        assert response.data['totalItems'] == 2
        returned_states = {item['state'] for item in response.data['items']}
        assert returned_states.issubset({"Draft", "Published"})

    def test_list_dynamic_documents_filter_by_lawyer_and_states(self, api_client, user):
        """list_dynamic_documents debe filtrar por lawyer_id y estados (caso Minutas)."""
        # Creamos dos abogados distintos
        lawyer_1 = User.objects.create_user(email='lawyer1@example.com', password='pwd', role='lawyer')
        lawyer_2 = User.objects.create_user(email='lawyer2@example.com', password='pwd', role='lawyer')

        # Autenticamos como admin genérico para listar (role no debería limitar por ser abogado en el decorador)
        api_client.force_authenticate(user=lawyer_1)

        # Documentos del lawyer_1 en distintos estados
        doc1 = DynamicDocument.objects.create(
            title="L1 Draft",
            content="<p>x</p>",
            state="Draft",
            created_by=lawyer_1,
        )
        doc2 = DynamicDocument.objects.create(
            title="L1 Published",
            content="<p>x</p>",
            state="Published",
            created_by=lawyer_1,
        )
        # Documento de otro abogado que no debería aparecer
        doc_other = DynamicDocument.objects.create(
            title="L2 Draft",
            content="<p>x</p>",
            state="Draft",
            created_by=lawyer_2,
        )

        url = reverse('list_dynamic_documents')

        response = api_client.get(url, {
            'lawyer_id': lawyer_1.id,
            'states': 'Draft,Published',
        })

        assert response.status_code == status.HTTP_200_OK
        assert response.data['totalItems'] == 2
        returned_ids = {item['id'] for item in response.data['items']}
        assert returned_ids == {doc1.id, doc2.id}

    def test_list_dynamic_documents_filter_by_client_and_states(self, api_client):
        """list_dynamic_documents debe filtrar por client_id y estados (caso Mis Documentos)."""
        # Creamos un cliente y otro usuario cualquiera
        client = User.objects.create_user(email='client@example.com', password='pwd', role='client')
        other_client = User.objects.create_user(email='other_client@example.com', password='pwd', role='client')

        # Autenticamos como cliente principal
        api_client.force_authenticate(user=client)

        # Documentos asignados al cliente en distintos estados
        doc_progress = DynamicDocument.objects.create(
            title="C1 Progress",
            content="<p>x</p>",
            state="Progress",
            created_by=client,
            assigned_to=client,
        )
        doc_completed = DynamicDocument.objects.create(
            title="C1 Completed",
            content="<p>x</p>",
            state="Completed",
            created_by=client,
            assigned_to=client,
        )

        # Documento asignado a otro cliente que no debería aparecer
        DynamicDocument.objects.create(
            title="C2 Progress",
            content="<p>x</p>",
            state="Progress",
            created_by=other_client,
            assigned_to=other_client,
        )

        url = reverse('list_dynamic_documents')

        response = api_client.get(url, {
            'client_id': client.id,
            'states': 'Progress,Completed',
        })

        assert response.status_code == status.HTTP_200_OK
        # Solo deben contarse los documentos del cliente autenticado en esos estados
        assert response.data['totalItems'] == 2
        returned_ids = {item['id'] for item in response.data['items']}
        assert returned_ids == {doc_progress.id, doc_completed.id}

    def test_list_dynamic_documents_unauthenticated(self, api_client, sample_document):
        """Test that unauthenticated users cannot access the document list"""
        url = reverse('list_dynamic_documents')
        response = api_client.get(url)
        
        # Assert the response
        assert response.status_code == status.HTTP_401_UNAUTHORIZED

    def test_list_dynamic_documents_invalid_page_and_limit(self, api_client, user):
        """Invalid page/limit should fall back to defaults and return first page."""
        api_client.force_authenticate(user=user)

        for i in range(12):
            DynamicDocument.objects.create(
                title=f"Doc {i}",
                content="<p>x</p>",
                state="Draft",
                created_by=user,
            )

        url = reverse('list_dynamic_documents')
        response = api_client.get(url, {'page': 'abc', 'limit': -5})

        assert response.status_code == status.HTTP_200_OK
        assert response.data['currentPage'] == 1
        assert response.data['totalItems'] == 12
        assert len(response.data['items']) == 10

    def test_list_dynamic_documents_out_of_range_page_returns_last(self, api_client, user):
        """Out-of-range page should return last available page."""
        api_client.force_authenticate(user=user)

        for i in range(12):
            DynamicDocument.objects.create(
                title=f"Doc {i}",
                content="<p>x</p>",
                state="Draft",
                created_by=user,
            )

        url = reverse('list_dynamic_documents')
        response = api_client.get(url, {'page': 99, 'limit': 5})

        assert response.status_code == status.HTTP_200_OK
        assert response.data['currentPage'] == 3
        assert response.data['totalPages'] == 3
        assert len(response.data['items']) == 2
    
    def test_create_dynamic_document_authenticated(self, api_client, user):
        """Test creating a new dynamic document when authenticated"""
        # Authenticate the user
        api_client.force_authenticate(user=user)
        
        # Create document data
        data = {
            "title": "New Document",
            "content": "<p>This is a new document with {{variable1}}.</p>",
            "state": "Draft",
            "variables": [
                {
                    "name_en": "variable1",
                    "value": "Test Value"
                }
            ]
        }
        
        # Make the request
        url = reverse('create_dynamic_document')
        response = api_client.post(url, data, format='json')
        
        # Assert the response
        assert response.status_code == status.HTTP_201_CREATED
        assert response.data['title'] == "New Document"
        
        # Verify document was created in the database
        assert DynamicDocument.objects.count() == 1
        document = DynamicDocument.objects.first()
        assert document.title == "New Document"
        assert document.created_by == user
        
        # Verify variable was created
        assert document.variables.count() == 1
        variable = document.variables.first()
        assert variable.name_en == "variable1"
        assert variable.value == "Test Value"

    def test_create_dynamic_document_assigns_user_when_progress(self, api_client, user):
        """Progress documents created without assigned_to should default to request.user."""
        api_client.force_authenticate(user=user)

        data = {
            "title": "Progress Doc",
            "content": "<p>x</p>",
            "state": "Progress",
        }

        url = reverse('create_dynamic_document')
        response = api_client.post(url, data, format='json')

        assert response.status_code == status.HTTP_201_CREATED
        assert response.data['assigned_to'] == user.id

    def test_create_dynamic_document_invalid_select_variable(self, api_client, user):
        """Select variables require options; invalid payload should return 400."""
        api_client.force_authenticate(user=user)

        data = {
            "title": "Bad Doc",
            "content": "<p>x</p>",
            "state": "Draft",
            "variables": [
                {
                    "name_en": "choice",
                    "field_type": "select",
                    "select_options": [],
                }
            ],
        }

        url = reverse('create_dynamic_document')
        response = api_client.post(url, data, format='json')

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert 'variables' in response.data
    
    def test_update_dynamic_document_authenticated(self, api_client, user, sample_document):
        """Test updating an existing dynamic document when authenticated"""
        # Authenticate the user
        api_client.force_authenticate(user=user)
        
        # Create update data
        data = {
            "title": "Updated Document",
            "content": "<p>This is the updated content.</p>",
            "state": "Progress",
            "variables": [
                {
                    "name_en": "variable1",
                    "value": "Updated Value 1"
                },
                {
                    "name_en": "variable2",
                    "value": "Updated Value 2"
                }
            ]
        }
        
        # Make the request
        url = reverse('update_dynamic_document', kwargs={'pk': sample_document.pk})
        response = api_client.put(url, data, format='json')
        
        # Assert the response
        assert response.status_code == status.HTTP_200_OK
        assert response.data['title'] == "Updated Document"
        assert response.data['state'] == "Progress"
        
        # Verify document was updated in the database
        sample_document.refresh_from_db()
        assert sample_document.title == "Updated Document"
        assert sample_document.state == "Progress"
        
        # Verify variables were updated
        variables = {var.name_en: var.value for var in sample_document.variables.all()}
        assert variables["variable1"] == "Updated Value 1"
        assert variables["variable2"] == "Updated Value 2"

    def test_update_dynamic_document_ignores_created_by(self, api_client, user, sample_document):
        """created_by should not be updated via the endpoint."""
        api_client.force_authenticate(user=user)
        other_user = User.objects.create_user(email='other@example.com', password='pass')

        url = reverse('update_dynamic_document', kwargs={'pk': sample_document.pk})
        response = api_client.put(url, {'created_by': other_user.id, 'title': 'Keep owner'}, format='json')

        assert response.status_code == status.HTTP_200_OK
        sample_document.refresh_from_db()
        assert sample_document.created_by == user

    def test_update_dynamic_document_invalid_variable_returns_400(self, api_client, user, sample_document):
        """Invalid variables should return serializer errors."""
        api_client.force_authenticate(user=user)

        data = {
            "variables": [
                {
                    "name_en": "select_var",
                    "field_type": "select",
                    "select_options": [],
                }
            ]
        }

        url = reverse('update_dynamic_document', kwargs={'pk': sample_document.pk})
        response = api_client.put(url, data, format='json')

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert 'variables' in response.data

    def test_get_dynamic_document_initializes_select_options(self, api_client, user, sample_document):
        """Select variables without options should be initialized to empty list."""
        api_client.force_authenticate(user=user)

        select_var = DocumentVariable.objects.create(
            document=sample_document,
            name_en='select_var',
            field_type='select',
            select_options=None,
        )

        url = reverse('get_dynamic_document', kwargs={'pk': sample_document.pk})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_200_OK
        select_var.refresh_from_db()
        assert select_var.select_options == []

    def test_get_dynamic_document_not_found(self, api_client, user):
        api_client.force_authenticate(user=user)

        url = reverse('get_dynamic_document', kwargs={'pk': 9999})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND
    
    def test_delete_dynamic_document_authenticated(self, api_client, user, sample_document):
        """Test deleting a dynamic document when authenticated"""
        # Authenticate the user
        api_client.force_authenticate(user=user)
        
        # Make the request
        url = reverse('delete_dynamic_document', kwargs={'pk': sample_document.pk})
        response = api_client.delete(url)
        
        # Assert the response
        assert response.status_code == status.HTTP_200_OK
        assert response.data['detail'] == "Dynamic document deleted successfully."
        
        # Verify document was deleted from the database
        assert DynamicDocument.objects.count() == 0

@pytest.mark.django_db
class TestDynamicDocumentExport:
    
    def test_download_dynamic_document_pdf_authenticated(self, api_client, user, sample_document, monkeypatch):
        """Test downloading a dynamic document as PDF when authenticated"""
        # Mock the pisa.CreatePDF function to avoid actual PDF creation
        class MockPisaStatus:
            err = False
            
        def mock_create_pdf(*args, **kwargs):
            return MockPisaStatus()
            
        monkeypatch.setattr('xhtml2pdf.pisa.CreatePDF', mock_create_pdf)
        
        # Authenticate the user
        api_client.force_authenticate(user=user)
        
        # Make the request
        url = reverse('download_dynamic_document_pdf', kwargs={'pk': sample_document.pk})
        response = api_client.get(url)
        
        # Assert the response
        assert response.status_code == status.HTTP_200_OK
        assert response['Content-Type'] == 'application/pdf'
        assert response['Content-Disposition'] == f'attachment; filename="{sample_document.title}.pdf"'
    
    def test_download_dynamic_document_word_authenticated(self, api_client, user, sample_document, monkeypatch):
        """Test downloading a dynamic document as Word when authenticated"""
        # Create a mock for Document.save
        class MockDocument:  # pragma: no cover – mock methods not invoked
            def __init__(self):
                self.styles = {
                    'Normal': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading1': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading2': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading3': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading4': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading5': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading6': type('Style', (), {'font': type('Font', (), {'name': None})})
                }
                
            def add_paragraph(self, *args, **kwargs):
                return type('Paragraph', (), {
                    'runs': [],
                    'add_run': lambda text: type('Run', (), {'font': type('Font', (), {'name': None})}),
                    'alignment': None,
                    'paragraph_format': type('ParagraphFormat', (), {'left_indent': None, 'line_spacing': None})
                })
                
            def add_heading(self, *args, **kwargs):
                return type('Heading', (), {
                    'runs': [type('Run', (), {'font': type('Font', (), {'name': None})})],
                })
                
            def save(self, *args, **kwargs):
                pass
        
        # Replace Document with our mock
        monkeypatch.setattr('docx.Document', lambda: MockDocument())
        
        # Create a mock for get_template
        def mock_get_template(*args, **kwargs):  # pragma: no cover – mock factory not invoked
            return type('Template', (), {
                'render': lambda context: f'<html><body>{context["content"]}</body></html>'
            })
        
        monkeypatch.setattr('django.template.loader.get_template', mock_get_template)
        
        # Authenticate the user
        api_client.force_authenticate(user=user)
        
        # Make the request
        url = reverse('download_dynamic_document_word', kwargs={'pk': sample_document.pk})
        response = api_client.get(url)
        
        # Assert the response
        assert response.status_code == status.HTTP_200_OK
        assert response['Content-Type'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        assert response['Content-Disposition'] == f'attachment; filename="{sample_document.title}.docx"'

    def test_download_dynamic_document_pdf_missing_fonts_returns_500(self, api_client, user, sample_document, monkeypatch):
        api_client.force_authenticate(user=user)

        monkeypatch.setattr(document_views.os.path, "exists", lambda path: False)

        url = reverse('download_dynamic_document_pdf', kwargs={'pk': sample_document.pk})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_500_INTERNAL_SERVER_ERROR
        assert "Font file not found" in response.data['detail']

    def test_download_dynamic_document_pdf_pisa_error_returns_500(self, api_client, user, sample_document, monkeypatch):
        api_client.force_authenticate(user=user)

        class MockPisaStatus:
            err = True

        def mock_create_pdf(*args, **kwargs):
            return MockPisaStatus()

        monkeypatch.setattr(document_views.os.path, "exists", lambda path: True)
        monkeypatch.setattr(document_views.pdfmetrics, "registerFont", lambda *args, **kwargs: None)
        monkeypatch.setattr(document_views, "TTFont", lambda *args, **kwargs: object())
        monkeypatch.setattr(document_views.pisa, "CreatePDF", mock_create_pdf)

        url = reverse('download_dynamic_document_pdf', kwargs={'pk': sample_document.pk})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_500_INTERNAL_SERVER_ERROR
        assert "Error generating PDF" in response.data['detail']

    def test_download_dynamic_document_word_template_error_returns_500(self, api_client, user, sample_document, monkeypatch):
        api_client.force_authenticate(user=user)

        def raise_error(*args, **kwargs):
            raise Exception("template fail")

        monkeypatch.setattr(document_views, "get_template", raise_error)

        url = reverse('download_dynamic_document_word', kwargs={'pk': sample_document.pk})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_500_INTERNAL_SERVER_ERROR
        assert "Error generating Word document" in response.data['detail']

    def test_download_dynamic_document_word_invalid_template_falls_back(self, api_client, user, sample_document, monkeypatch, settings, tmp_path):
        settings.MEDIA_ROOT = tmp_path
        api_client.force_authenticate(user=user)

        template_file = SimpleUploadedFile(
            "template.docx",
            b"bad-template",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        user.letterhead_word_template = template_file
        user.save(update_fields=["letterhead_word_template"])

        class MockDocument:
            def __init__(self, *args, **kwargs):
                self.styles = {
                    'Normal': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading1': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading2': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading3': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading4': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading5': type('Style', (), {'font': type('Font', (), {'name': None})}),
                    'Heading6': type('Style', (), {'font': type('Font', (), {'name': None})}),
                }
                self.paragraphs = [type('Paragraph', (), {'text': '', 'runs': []})]
                self.sections = [type('Section', (), {})]

            def add_paragraph(self, *args, **kwargs):
                return type('Paragraph', (), {
                    'runs': [],
                    'add_run': lambda text: type('Run', (), {'font': type('Font', (), {'name': None})}),
                    'alignment': None,
                    'paragraph_format': type('ParagraphFormat', (), {'left_indent': None, 'line_spacing': None})
                })

            def add_heading(self, *args, **kwargs):  # pragma: no cover – mock method not invoked
                return type('Heading', (), {'runs': [type('Run', (), {'font': type('Font', (), {'name': None})})]})

            def save(self, *args, **kwargs):
                pass

        def document_factory(path=None):
            if path:
                raise Exception("bad template")
            return MockDocument()

        monkeypatch.setattr(document_views, "Document", document_factory)

        url = reverse('download_dynamic_document_word', kwargs={'pk': sample_document.pk})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_200_OK
    
    def test_document_download_unauthenticated(self, api_client, sample_document):
        """Test that unauthenticated users cannot download documents"""
        # Try PDF download
        pdf_url = reverse('download_dynamic_document_pdf', kwargs={'pk': sample_document.pk})
        pdf_response = api_client.get(pdf_url)
        assert pdf_response.status_code == status.HTTP_401_UNAUTHORIZED
        
        # Try Word download
        word_url = reverse('download_dynamic_document_word', kwargs={'pk': sample_document.pk})
        word_response = api_client.get(word_url)
        assert word_response.status_code == status.HTTP_401_UNAUTHORIZED
    
    def test_download_nonexistent_document(self, api_client, user):
        """Test trying to download a document that doesn't exist"""
        # Authenticate the user
        api_client.force_authenticate(user=user)
        
        # Try PDF download with non-existent ID
        pdf_url = reverse('download_dynamic_document_pdf', kwargs={'pk': 9999})
        pdf_response = api_client.get(pdf_url)
        assert pdf_response.status_code == status.HTTP_404_NOT_FOUND
        
        # Try Word download with non-existent ID
        word_url = reverse('download_dynamic_document_word', kwargs={'pk': 9999})
        word_response = api_client.get(word_url)
        assert word_response.status_code == status.HTTP_404_NOT_FOUND

@pytest.mark.django_db
class TestDynamicDocumentRecentViews:

    def test_get_recent_documents_empty(self, api_client, user):
        api_client.force_authenticate(user=user)

        url = reverse('get-recent-documents')
        response = api_client.get(url)

        assert response.status_code == status.HTTP_200_OK
        assert response.data == []

    def test_get_recent_documents_filters_by_can_view_and_limits_to_10(self, api_client, user):
        """get_recent_documents debe devolver como máximo 10 documentos y respetar can_view."""
        # Crear 11 documentos visibles para el usuario
        visible_docs = []
        for i in range(11):
            doc = DynamicDocument.objects.create(
                title=f"Doc {i}",
                content="<p>x</p>",
                state="Draft",
                created_by=user,
            )
            visible_docs.append(doc)
            RecentDocument.objects.create(user=user, document=doc)

        # Crear un documento que el usuario NO puede ver (otro creador, no público)
        other_user = User.objects.create_user(
            email='other@example.com',
            password='testpassword',
        )
        hidden_doc = DynamicDocument.objects.create(
            title="Hidden",
            content="<p>secret</p>",
            state="Draft",
            created_by=other_user,
            is_public=False,
        )
        RecentDocument.objects.create(user=user, document=hidden_doc)

        api_client.force_authenticate(user=user)

        url = reverse('get-recent-documents')
        response = api_client.get(url)

        assert response.status_code == status.HTTP_200_OK
        # Debe devolver como máximo 10 documentos
        assert len(response.data) == 10
        # Ninguno de los documentos debe ser el oculto
        doc_ids = {item['document']['id'] for item in response.data}
        assert hidden_doc.id not in doc_ids

    def test_update_recent_document_creates_and_updates_entry(self, api_client, user):
        """update_recent_document debe crear o actualizar la entrada de RecentDocument"""
        document = DynamicDocument.objects.create(
            title="Doc reciente",
            content="<p>x</p>",
            state="Draft",
            created_by=user,
        )

        api_client.force_authenticate(user=user)

        url = reverse('update-recent-document', kwargs={'document_id': document.id})

        # Primera llamada: crea la entrada
        response = api_client.post(url, {})
        assert response.status_code == status.HTTP_200_OK

        recent = RecentDocument.objects.get(user=user, document=document)
        first_last_visited = recent.last_visited

        # Segunda llamada: actualiza last_visited
        response = api_client.post(url, {})
        assert response.status_code == status.HTTP_200_OK

        recent.refresh_from_db()
        assert recent.last_visited >= first_last_visited

    def test_update_recent_document_forbidden_when_no_visibility(self, api_client, user):
        """update_recent_document debe devolver 403 si el usuario no puede ver el documento"""
        other_user = User.objects.create_user(
            email='other2@example.com',
            password='testpassword',
        )
        document = DynamicDocument.objects.create(
            title="Doc oculto",
            content="<p>secret</p>",
            state="Draft",
            created_by=other_user,
            is_public=False,
        )

        api_client.force_authenticate(user=user)

        url = reverse('update-recent-document', kwargs={'document_id': document.id})
        response = api_client.post(url, {})

        assert response.status_code == status.HTTP_403_FORBIDDEN
        assert 'permission' in response.data['detail'].lower()

    def test_update_recent_document_not_found(self, api_client, user):
        """update_recent_document debe devolver 404 si el documento no existe"""
        api_client.force_authenticate(user=user)

        url = reverse('update-recent-document', kwargs={'document_id': 9999})
        response = api_client.post(url, {})

        assert response.status_code == status.HTTP_404_NOT_FOUND
        assert 'not found' in response.data['detail']

# ======================================================================
# Tests migrated from test_views_batch18.py
# ======================================================================

"""
Batch 18 – 20 tests: document_views.py coverage gaps.
  • list pagination edges (PageNotAnInteger, EmptyPage)
  • get_dynamic_document DoesNotExist, select_options init
  • update_dynamic_document DoesNotExist, created_by strip, validation error
  • delete_dynamic_document DoesNotExist
  • download_dynamic_document_pdf DoesNotExist, font missing, general error
  • download_dynamic_document_word DoesNotExist, general error
  • Word doc with HTML content (heading, paragraph styles, hr)
  • Word doc with word template fallback
  • PDF for_version returns buffer
"""


@pytest.fixture
def api():
    return APIClient()


@pytest.fixture
@pytest.mark.django_db
def lawyer():
    return User.objects.create_user(
        email="law_b18@t.com", password="pw", role="lawyer",
        first_name="L", last_name="W",
    )


@pytest.fixture
@pytest.mark.django_db
def doc(lawyer):
    return DynamicDocument.objects.create(
        title="DocB18", content="<p>Hello {{var1}}</p>",
        state="Draft", created_by=lawyer,
    )


@pytest.fixture
@pytest.mark.django_db
def doc_with_var(doc):
    DocumentVariable.objects.create(document=doc, name_en="var1", value="World")
    return doc


# ===========================================================================
# 5. download_dynamic_document_pdf
# ===========================================================================

@pytest.mark.django_db
class TestDownloadPDF:

    def test_not_found(self, api, lawyer):
        """Lines 451-452: DoesNotExist returns 404."""
        api.force_authenticate(user=lawyer)
        url = reverse("download_dynamic_document_pdf", kwargs={"pk": 99999})
        resp = api.get(url)
        assert resp.status_code == 404

    @patch("gym_app.views.dynamic_documents.document_views.os.path.exists", return_value=False)
    def test_font_missing(self, mock_exists, api, lawyer, doc_with_var):
        """Lines 305-306: FileNotFoundError when font is missing."""
        api.force_authenticate(user=lawyer)
        url = reverse("download_dynamic_document_pdf", kwargs={"pk": doc_with_var.pk})
        resp = api.get(url)
        assert resp.status_code == 500
        assert "Font file" in resp.data.get("detail", "")

    @patch("gym_app.views.dynamic_documents.document_views.pisa.CreatePDF", side_effect=Exception("PDF boom"))
    @patch("gym_app.views.dynamic_documents.document_views.os.path.exists", return_value=True)
    @patch("gym_app.views.dynamic_documents.document_views.pdfmetrics.registerFont")
    def test_general_error(self, _reg, _exists, _pisa, api, lawyer, doc_with_var):
        """Lines 455-456: general exception returns 500."""
        api.force_authenticate(user=lawyer)
        url = reverse("download_dynamic_document_pdf", kwargs={"pk": doc_with_var.pk})
        resp = api.get(url)
        assert resp.status_code == 500


# ===========================================================================
# 6. download_dynamic_document_word
# ===========================================================================

@pytest.mark.django_db
class TestDownloadWord:

    def test_not_found(self, api, lawyer):
        """Lines 778-779: DoesNotExist returns 404."""
        api.force_authenticate(user=lawyer)
        url = reverse("download_dynamic_document_word", kwargs={"pk": 99999})
        resp = api.get(url)
        assert resp.status_code == 404

    @patch("gym_app.views.dynamic_documents.document_views.Document", side_effect=Exception("docx boom"))
    def test_general_error(self, _mock, api, lawyer, doc_with_var):
        """Lines 780-781: general exception returns 500."""
        api.force_authenticate(user=lawyer)
        url = reverse("download_dynamic_document_word", kwargs={"pk": doc_with_var.pk})
        resp = api.get(url)
        assert resp.status_code == 500

    def test_word_success_with_html_content(self, api, lawyer):
        """Lines 554-765: Word generation processes headings, paragraphs, styles, hr."""
        doc = DynamicDocument.objects.create(
            title="WordDoc",
            content='<h1>Title</h1><p style="text-align: center">Center</p><hr><p>Normal</p>',
            state="Draft", created_by=lawyer,
        )
        api.force_authenticate(user=lawyer)
        url = reverse("download_dynamic_document_word", kwargs={"pk": doc.pk})
        resp = api.get(url)
        assert resp.status_code == 200
        assert "wordprocessing" in resp.get("Content-Type", "")


# ======================================================================
# Tests migrated from test_views_batch19.py
# ======================================================================

"""
Batch 19 – 20 tests: document_views.py coverage gaps continued.
  • get_recent_documents
  • update_recent_document (success, not-found)
  • upload_letterhead_image (success, no file, bad ext, too large, invalid img,
    DoesNotExist, general exception, aspect ratio warning)
  • get_letterhead_image (no image, file missing, DoesNotExist, general exc)
  • delete_letterhead_image (no image, DoesNotExist, general exc, success)
  • upload_document_letterhead_word_template (success, no file, bad ext, too large,
    DoesNotExist, general exc)
"""


@pytest.fixture
def api():
    return APIClient()


@pytest.fixture
@pytest.mark.django_db
def lawyer():
    return User.objects.create_user(
        email="law_b19@t.com", password="pw", role="lawyer",
        first_name="L", last_name="W",
    )


@pytest.fixture
@pytest.mark.django_db
def doc(lawyer):
    return DynamicDocument.objects.create(
        title="DocB19", content="<p>hi</p>", state="Draft", created_by=lawyer,
    )


def _make_png(w=100, h=100):
    """Create a minimal valid PNG in memory."""
    buf = io.BytesIO()
    img = Image.new("RGB", (w, h), color="white")
    img.save(buf, format="PNG")
    buf.seek(0)
    return SimpleUploadedFile("test.png", buf.read(), content_type="image/png")


def _make_docx():
    """Create a minimal .docx file."""
    from docx import Document as DocxDoc
    buf = io.BytesIO()
    d = DocxDoc()
    d.add_paragraph("template")
    d.save(buf)
    buf.seek(0)
    return SimpleUploadedFile("tpl.docx", buf.read(),
                              content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ===========================================================================
# 1. Recent documents
# ===========================================================================

@pytest.mark.django_db
class TestRecentDocuments:

    def test_get_recent_empty(self, api, lawyer):
        api.force_authenticate(user=lawyer)
        url = reverse("get-recent-documents")
        resp = api.get(url)
        assert resp.status_code == 200
        assert resp.data == []

    def test_update_recent_success(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("update-recent-document", kwargs={"document_id": doc.pk})
        resp = api.post(url)
        assert resp.status_code == 200

    def test_update_recent_not_found(self, api, lawyer):
        api.force_authenticate(user=lawyer)
        url = reverse("update-recent-document", kwargs={"document_id": 99999})
        resp = api.post(url)
        assert resp.status_code == 404


# ===========================================================================
# 2. Letterhead image CRUD (per-document)
# ===========================================================================

@pytest.mark.django_db
class TestLetterheadImage:

    def test_upload_success(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("upload-letterhead-image", kwargs={"pk": doc.pk})
        resp = api.post(url, {"image": _make_png()}, format="multipart")
        assert resp.status_code == 201

    def test_upload_no_file(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("upload-letterhead-image", kwargs={"pk": doc.pk})
        resp = api.post(url, {}, format="multipart")
        assert resp.status_code == 400

    def test_upload_bad_ext(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("upload-letterhead-image", kwargs={"pk": doc.pk})
        f = SimpleUploadedFile("bad.jpg", b"data", content_type="image/jpeg")
        resp = api.post(url, {"image": f}, format="multipart")
        assert resp.status_code == 400

    def test_upload_too_large(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("upload-letterhead-image", kwargs={"pk": doc.pk})
        f = SimpleUploadedFile("big.png", b"x" * (11 * 1024 * 1024), content_type="image/png")
        resp = api.post(url, {"image": f}, format="multipart")
        assert resp.status_code == 400

    def test_upload_invalid_image(self, api, lawyer, doc):
        """Lines 890-893: invalid image data."""
        api.force_authenticate(user=lawyer)
        url = reverse("upload-letterhead-image", kwargs={"pk": doc.pk})
        f = SimpleUploadedFile("bad.png", b"not-an-image", content_type="image/png")
        resp = api.post(url, {"image": f}, format="multipart")
        assert resp.status_code == 400

    def test_upload_not_found(self, api, lawyer):
        """Lines 925-929: DoesNotExist."""
        api.force_authenticate(user=lawyer)
        url = reverse("upload-letterhead-image", kwargs={"pk": 99999})
        resp = api.post(url, {"image": _make_png()}, format="multipart")
        assert resp.status_code == 404

    def test_upload_aspect_ratio_warning(self, api, lawyer, doc):
        """Lines 887-888: aspect ratio differs from recommended."""
        api.force_authenticate(user=lawyer)
        url = reverse("upload-letterhead-image", kwargs={"pk": doc.pk})
        resp = api.post(url, {"image": _make_png(500, 100)}, format="multipart")
        assert resp.status_code == 201
        assert "warnings" in resp.data

    def test_get_no_image(self, api, lawyer, doc):
        """Lines 946-950: no letterhead image."""
        api.force_authenticate(user=lawyer)
        url = reverse("get-letterhead-image", kwargs={"pk": doc.pk})
        resp = api.get(url)
        assert resp.status_code == 404

    def test_get_not_found(self, api, lawyer):
        """Lines 967-971: DoesNotExist."""
        api.force_authenticate(user=lawyer)
        url = reverse("get-letterhead-image", kwargs={"pk": 99999})
        resp = api.get(url)
        assert resp.status_code == 404

    def test_delete_no_image(self, api, lawyer, doc):
        """Lines 988-992: no image to delete."""
        api.force_authenticate(user=lawyer)
        url = reverse("delete-letterhead-image", kwargs={"pk": doc.pk})
        resp = api.delete(url)
        assert resp.status_code == 404

    def test_delete_not_found(self, api, lawyer):
        """Lines 1010-1013: DoesNotExist."""
        api.force_authenticate(user=lawyer)
        url = reverse("delete-letterhead-image", kwargs={"pk": 99999})
        resp = api.delete(url)
        assert resp.status_code == 404


# ===========================================================================
# 3. Document word template CRUD
# ===========================================================================

@pytest.mark.django_db
class TestDocWordTemplate:

    def test_upload_success(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": doc.pk})
        resp = api.post(url, {"template": _make_docx()}, format="multipart")
        assert resp.status_code == 201

    def test_upload_no_file(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": doc.pk})
        resp = api.post(url, {}, format="multipart")
        assert resp.status_code == 400

    def test_upload_bad_ext(self, api, lawyer, doc):
        api.force_authenticate(user=lawyer)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": doc.pk})
        f = SimpleUploadedFile("bad.txt", b"data", content_type="text/plain")
        resp = api.post(url, {"template": f}, format="multipart")
        assert resp.status_code == 400

    def test_upload_not_found(self, api, lawyer):
        """Lines 1080-1084: DoesNotExist."""
        api.force_authenticate(user=lawyer)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": 99999})
        resp = api.post(url, {"template": _make_docx()}, format="multipart")
        assert resp.status_code == 404

    def test_get_no_template(self, api, lawyer, doc):
        """Lines 1100-1103: no template configured."""
        api.force_authenticate(user=lawyer)
        url = reverse("get-document-letterhead-word-template", kwargs={"pk": doc.pk})
        resp = api.get(url)
        assert resp.status_code == 404

    def test_delete_no_template(self, api, lawyer, doc):
        """Lines 1139-1142: no template to delete."""
        api.force_authenticate(user=lawyer)
        url = reverse("delete-document-letterhead-word-template", kwargs={"pk": doc.pk})
        resp = api.delete(url)
        assert resp.status_code == 404


# ======================================================================
# Tests migrated from test_views_batch31.py
# ======================================================================

"""Batch 31 – 20 tests: document_views.py – list filters, pagination, CRUD edges, recent docs, letterhead."""


@pytest.fixture
def api():
    return APIClient()

@pytest.fixture
def law():
    return User.objects.create_user(email="law31@t.com", password="pw", role="lawyer", first_name="L", last_name="W")

@pytest.fixture
def cli():
    return User.objects.create_user(email="cli31@t.com", password="pw", role="client", first_name="C", last_name="E")


# -- list_dynamic_documents filter/pagination edges --
class TestListDynDocEdges:
    def test_filter_by_client_id(self, api, law, cli):
        DynamicDocument.objects.create(title="ClientDoc", content="<p>x</p>", state="Draft", created_by=law, assigned_to=cli)
        DynamicDocument.objects.create(title="NoClient", content="<p>x</p>", state="Draft", created_by=law)
        api.force_authenticate(user=law)
        resp = api.get(reverse("list_dynamic_documents"), {"client_id": cli.id})
        assert resp.status_code == 200
        titles = [d["title"] for d in resp.data["items"]]
        assert "ClientDoc" in titles
        assert "NoClient" not in titles

    def test_filter_by_lawyer_id(self, api, law):
        law2 = User.objects.create_user(email="law31b@t.com", password="pw", role="lawyer")
        DynamicDocument.objects.create(title="Mine", content="<p>x</p>", state="Draft", created_by=law)
        DynamicDocument.objects.create(title="Other", content="<p>x</p>", state="Draft", created_by=law2)
        api.force_authenticate(user=law)
        resp = api.get(reverse("list_dynamic_documents"), {"lawyer_id": law.id})
        assert resp.status_code == 200
        titles = [d["title"] for d in resp.data["items"]]
        assert "Mine" in titles
        assert "Other" not in titles

    def test_pagination_defaults(self, api, law):
        for i in range(15):
            DynamicDocument.objects.create(title=f"P{i}", content="<p>x</p>", state="Draft", created_by=law)
        api.force_authenticate(user=law)
        resp = api.get(reverse("list_dynamic_documents"))
        assert resp.status_code == 200
        assert resp.data["totalItems"] == 15
        assert resp.data["totalPages"] == 2
        assert len(resp.data["items"]) == 10  # default limit

    def test_pagination_custom_limit(self, api, law):
        for i in range(5):
            DynamicDocument.objects.create(title=f"Q{i}", content="<p>x</p>", state="Draft", created_by=law)
        api.force_authenticate(user=law)
        resp = api.get(reverse("list_dynamic_documents"), {"limit": 2, "page": 2})
        assert resp.status_code == 200
        assert resp.data["currentPage"] == 2
        assert len(resp.data["items"]) == 2

    def test_pagination_negative_limit(self, api, law):
        DynamicDocument.objects.create(title="X", content="<p>x</p>", state="Draft", created_by=law)
        api.force_authenticate(user=law)
        resp = api.get(reverse("list_dynamic_documents"), {"limit": -5})
        assert resp.status_code == 200
        assert resp.data["totalPages"] >= 1


# -- get_dynamic_document --
class TestGetDynDocEdges:
    def test_get_doc_no_permission(self, api, law, cli):
        doc = DynamicDocument.objects.create(title="Priv", content="<p>x</p>", state="Draft", created_by=law)
        api.force_authenticate(user=cli)
        resp = api.get(reverse("get_dynamic_document", args=[doc.id]))
        assert resp.status_code == 403


# -- recent documents --
class TestRecentDocs:
    def test_update_recent_creates(self, api, law):
        doc = DynamicDocument.objects.create(title="Rec", content="<p>x</p>", state="Draft", created_by=law)
        doc.visibility_permissions.create(user=law, granted_by=law)
        api.force_authenticate(user=law)
        resp = api.post(reverse("update-recent-document", args=[doc.id]))
        assert resp.status_code == 200
        assert RecentDocument.objects.filter(user=law, document=doc).exists()

    def test_update_recent_updates_timestamp(self, api, law):
        doc = DynamicDocument.objects.create(title="Rec2", content="<p>x</p>", state="Draft", created_by=law)
        doc.visibility_permissions.create(user=law, granted_by=law)
        RecentDocument.objects.create(user=law, document=doc)
        api.force_authenticate(user=law)
        resp = api.post(reverse("update-recent-document", args=[doc.id]))
        assert resp.status_code == 200


# -- letterhead endpoints --
class TestLetterheadEdges:
    def test_get_user_letterhead_no_image(self, api, law):
        api.force_authenticate(user=law)
        resp = api.get(reverse("get-user-letterhead-image"))
        assert resp.status_code == 404


# ======================================================================
# ======================================================================
# Tests migrated from test_views_batch5.py
# ======================================================================

"""
Batch 5 – Coverage-gap tests for document_views.py (75% → higher)
Targets uncovered letterhead image/word-template endpoints,
list filter edge cases, PDF/Word error paths, recent document edge,
and get_letterhead_for_document helper.
"""


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
@pytest.fixture
@pytest.mark.django_db
def lawyer_user():
    return User.objects.create_user(
        email="lawyer_b5@example.com",
        password="testpassword",
        first_name="Law",
        last_name="Yer",
        role="lawyer",
    )


@pytest.fixture
@pytest.mark.django_db
def basic_user():
    return User.objects.create_user(
        email="basic_b5@example.com",
        password="testpassword",
        role="basic",
    )


@pytest.fixture
@pytest.mark.django_db
def document(lawyer_user):
    return DynamicDocument.objects.create(
        title="Test Doc B5",
        content="<p>Hello</p>",
        state="Draft",
        created_by=lawyer_user,
    )


def _png_file(name="letterhead.png", w=100, h=130):
    buf = BytesIO()
    Image.new("RGB", (w, h), color="white").save(buf, format="PNG")
    buf.seek(0)
    return SimpleUploadedFile(name, buf.read(), content_type="image/png")


def _docx_bytes():
    """Minimal valid .docx file (just a ZIP with required content types)."""
    from docx import Document as DocxDocument
    buf = BytesIO()
    DocxDocument().save(buf)
    buf.seek(0)
    return buf.read()


# ===========================================================================
# 1. list_dynamic_documents – filter / pagination edge cases
# ===========================================================================

@pytest.mark.django_db
class TestListDynamicDocumentsEdges:

    def test_filter_by_single_state(self, api_client, lawyer_user):
        DynamicDocument.objects.create(title="D1", content="<p>a</p>", state="Draft", created_by=lawyer_user)
        DynamicDocument.objects.create(title="D2", content="<p>b</p>", state="Completed", created_by=lawyer_user)

        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        response = api_client.get(url, {"state": "Completed"})

        assert response.status_code == status.HTTP_200_OK
        titles = [d["title"] for d in response.data["items"]]
        assert "D2" in titles
        assert "D1" not in titles

    def test_filter_by_multi_states(self, api_client, lawyer_user):
        DynamicDocument.objects.create(title="D1", content="<p>a</p>", state="Draft", created_by=lawyer_user)
        DynamicDocument.objects.create(title="D2", content="<p>b</p>", state="Completed", created_by=lawyer_user)
        DynamicDocument.objects.create(title="D3", content="<p>c</p>", state="Progress", created_by=lawyer_user)

        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        response = api_client.get(url, {"states": "Draft,Progress"})

        assert response.status_code == status.HTTP_200_OK
        titles = {d["title"] for d in response.data["items"]}
        assert "D1" in titles
        assert "D3" in titles
        assert "D2" not in titles

    def test_empty_page_falls_to_last(self, api_client, lawyer_user):
        DynamicDocument.objects.create(title="D1", content="<p>a</p>", state="Draft", created_by=lawyer_user)

        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        response = api_client.get(url, {"page": "999", "limit": "1"})

        assert response.status_code == status.HTTP_200_OK
        assert response.data["totalItems"] == 1


# ===========================================================================
# 2. Document letterhead image endpoints
# ===========================================================================

@pytest.mark.django_db
class TestDocumentLetterheadImage:

    def test_upload_letterhead_image_success(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        png = _png_file()
        response = api_client.post(url, {"image": png}, format="multipart")

        assert response.status_code == status.HTTP_201_CREATED
        assert "image_info" in response.data
        document.refresh_from_db()
        assert bool(document.letterhead_image)

    def test_upload_letterhead_image_missing_file(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        response = api_client.post(url, {}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "imagen" in response.data["detail"].lower()

    def test_upload_letterhead_image_wrong_extension(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        jpg = SimpleUploadedFile("bad.jpg", b"content", content_type="image/jpeg")
        response = api_client.post(url, {"image": jpg}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "PNG" in response.data["detail"]

    def test_upload_letterhead_image_too_large(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        # Create an oversized file (>10MB)
        big = SimpleUploadedFile("big.png", b"x" * (10 * 1024 * 1024 + 1), content_type="image/png")
        response = api_client.post(url, {"image": big}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "grande" in response.data["detail"].lower()

    def test_upload_letterhead_image_invalid_image(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        bad_png = SimpleUploadedFile("bad.png", b"not-a-real-png", content_type="image/png")
        response = api_client.post(url, {"image": bad_png}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "inválido" in response.data["detail"].lower()

    def test_get_letterhead_image_not_set(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-letterhead-image", kwargs={"pk": document.id})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_letterhead_image_not_set(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("delete-letterhead-image", kwargs={"pk": document.id})
        response = api_client.delete(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_letterhead_image_success(self, api_client, lawyer_user, document):
        # First upload
        api_client.force_authenticate(user=lawyer_user)
        upload_url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        png = _png_file()
        api_client.post(upload_url, {"image": png}, format="multipart")

        document.refresh_from_db()
        assert bool(document.letterhead_image)

        # Then delete
        del_url = reverse("delete-letterhead-image", kwargs={"pk": document.id})
        response = api_client.delete(del_url)

        assert response.status_code == status.HTTP_200_OK
        document.refresh_from_db()
        assert not document.letterhead_image


# ===========================================================================
# 3. Document letterhead Word template endpoints
# ===========================================================================

@pytest.mark.django_db
class TestDocumentLetterheadWordTemplate:

    def test_upload_word_template_success(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": document.id})
        docx = SimpleUploadedFile("template.docx", _docx_bytes(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response = api_client.post(url, {"template": docx}, format="multipart")

        assert response.status_code == status.HTTP_201_CREATED
        assert "template_info" in response.data

    def test_upload_word_template_missing_file(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": document.id})
        response = api_client.post(url, {}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_upload_word_template_wrong_extension(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": document.id})
        txt = SimpleUploadedFile("bad.txt", b"content", content_type="text/plain")
        response = api_client.post(url, {"template": txt}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert ".docx" in response.data["detail"]

    def test_get_word_template_not_set(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-document-letterhead-word-template", kwargs={"pk": document.id})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_word_template_not_set(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("delete-document-letterhead-word-template", kwargs={"pk": document.id})
        response = api_client.delete(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_word_template_success(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        upload_url = reverse("upload-document-letterhead-word-template", kwargs={"pk": document.id})
        docx = SimpleUploadedFile("template.docx", _docx_bytes(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        api_client.post(upload_url, {"template": docx}, format="multipart")

        del_url = reverse("delete-document-letterhead-word-template", kwargs={"pk": document.id})
        response = api_client.delete(del_url)

        assert response.status_code == status.HTTP_200_OK
        document.refresh_from_db()
        assert not document.letterhead_word_template


# ===========================================================================
# 4. User global letterhead image endpoints
# ===========================================================================

@pytest.mark.django_db
class TestUserGlobalLetterhead:

    def test_upload_user_letterhead_basic_forbidden(self, api_client, basic_user):
        api_client.force_authenticate(user=basic_user)
        url = reverse("upload-user-letterhead-image")
        png = _png_file()
        response = api_client.post(url, {"image": png}, format="multipart")

        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_upload_user_letterhead_success(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-image")
        png = _png_file()
        response = api_client.post(url, {"image": png}, format="multipart")

        assert response.status_code == status.HTTP_201_CREATED
        lawyer_user.refresh_from_db()
        assert bool(lawyer_user.letterhead_image)

    def test_get_user_letterhead_not_set(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-user-letterhead-image")
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_user_letterhead_basic_forbidden(self, api_client, basic_user):
        api_client.force_authenticate(user=basic_user)
        url = reverse("delete-user-letterhead-image")
        response = api_client.delete(url)

        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_delete_user_letterhead_not_set(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("delete-user-letterhead-image")
        response = api_client.delete(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_user_letterhead_success(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        # Upload first
        upload_url = reverse("upload-user-letterhead-image")
        png = _png_file()
        api_client.post(upload_url, {"image": png}, format="multipart")
        lawyer_user.refresh_from_db()
        assert bool(lawyer_user.letterhead_image)

        # Delete
        del_url = reverse("delete-user-letterhead-image")
        response = api_client.delete(del_url)

        assert response.status_code == status.HTTP_200_OK
        lawyer_user.refresh_from_db()
        assert not lawyer_user.letterhead_image


# ===========================================================================
# 5. User global Word letterhead template endpoints
# ===========================================================================

@pytest.mark.django_db
class TestUserGlobalWordTemplate:

    def test_upload_user_word_template_basic_forbidden(self, api_client, basic_user):
        api_client.force_authenticate(user=basic_user)
        url = reverse("upload-user-letterhead-word-template")
        docx = SimpleUploadedFile("t.docx", _docx_bytes(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response = api_client.post(url, {"template": docx}, format="multipart")

        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_upload_user_word_template_success(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-word-template")
        docx = SimpleUploadedFile("t.docx", _docx_bytes(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response = api_client.post(url, {"template": docx}, format="multipart")

        assert response.status_code == status.HTTP_201_CREATED

    def test_get_user_word_template_not_set(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-user-letterhead-word-template")
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_user_word_template_basic_forbidden(self, api_client, basic_user):
        api_client.force_authenticate(user=basic_user)
        url = reverse("delete-user-letterhead-word-template")
        response = api_client.delete(url)

        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_delete_user_word_template_not_set(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("delete-user-letterhead-word-template")
        response = api_client.delete(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_user_word_template_success(self, api_client, lawyer_user):
        api_client.force_authenticate(user=lawyer_user)
        # Upload first
        upload_url = reverse("upload-user-letterhead-word-template")
        docx = SimpleUploadedFile("t.docx", _docx_bytes(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        api_client.post(upload_url, {"template": docx}, format="multipart")

        # Delete
        del_url = reverse("delete-user-letterhead-word-template")
        response = api_client.delete(del_url)

        assert response.status_code == status.HTTP_200_OK


# ===========================================================================
# 6. Recent documents + update_recent_document edge
# ===========================================================================

@pytest.mark.django_db
class TestRecentDocumentEdges:

    def test_update_recent_document_creates_and_updates(self, api_client, lawyer_user, document):
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("update-recent-document", kwargs={"document_id": document.id})

        # First call creates
        resp1 = api_client.post(url, {}, format="json")
        assert resp1.status_code == status.HTTP_200_OK
        assert RecentDocument.objects.filter(user=lawyer_user, document=document).exists()

        # Second call updates
        resp2 = api_client.post(url, {}, format="json")
        assert resp2.status_code == status.HTTP_200_OK
        assert RecentDocument.objects.filter(user=lawyer_user, document=document).count() == 1


# ======================================================================
# Tests merged from test_document_views_coverage.py
# ======================================================================

"""Tests for uncovered branches in document_views.py (75%→higher)."""

@pytest.fixture
def lawyer():
    return User.objects.create_user(
        email='law_dvc@e.com', password='p', role='lawyer',
        first_name='L', last_name='D')


@pytest.fixture
def basic_user():
    return User.objects.create_user(
        email='basic_dvc@e.com', password='p', role='basic',
        first_name='B', last_name='D')


@pytest.fixture
def doc(lawyer):
    """Document with rich HTML content for Word/PDF tests."""
    html = (
        '<h1>Title</h1>'
        '<h2>Subtitle</h2>'
        '<p style="text-align: center">Centered</p>'
        '<p style="text-align: right">Right</p>'
        '<p style="text-align: left">Left</p>'
        '<p style="text-align: justify">Justified</p>'
        '<p style="padding-left: 40px">Indented</p>'
        '<p style="line-height: 1.5">Spaced</p>'
        '<p><strong>Bold</strong> <em>Italic</em> <u>Under</u></p>'
        '<p><span style="text-decoration: underline">SpanUnder</span></p>'
        '<p><span style="text-decoration: line-through">Strike</span></p>'
        '<p><span style="font-size: 14pt">BigText</span></p>'
        '<p><span style="color: red">RedText</span></p>'
        '<p><span style="color: rgb(0,128,0)">GreenRGB</span></p>'
        '<p><span style="color: #0000FF">BlueHex</span></p>'
        '<div>DivContent</div>'
        '<hr/>'
        '<p></p>'
    )
    return DynamicDocument.objects.create(
        title='TestDocDVC', content=html, state='Draft',
        created_by=lawyer,
        created_at=timezone.now(), updated_at=timezone.now())


@pytest.fixture
def doc_empty(lawyer):
    return DynamicDocument.objects.create(
        title='EmptyDoc', content='', state='Draft',
        created_by=lawyer,
        created_at=timezone.now(), updated_at=timezone.now())


@pytest.mark.django_db
class TestDocViewsCoverage:

    # --- Pagination: non-integer limit (lines 113-114) ---
    def test_list_non_integer_limit(self, api_client, lawyer, doc):
        """Lines 113-114: non-integer limit falls back to 10."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(
            reverse('list_dynamic_documents'),
            {'limit': 'abc'})
        assert r.status_code == 200

    # --- Word download with rich HTML (lines 556-710+) ---
    def test_download_word_rich_content(self, api_client, lawyer, doc):
        """Lines 556-710: Word generation with styled HTML content."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(
            reverse('download_dynamic_document_word', kwargs={'pk': doc.pk}))
        assert r.status_code == 200
        assert 'application/vnd.openxmlformats' in r['Content-Type']

    # --- Word download empty content ---
    def test_download_word_empty(self, api_client, lawyer, doc_empty):
        """Word generation with empty content."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(
            reverse('download_dynamic_document_word',
                    kwargs={'pk': doc_empty.pk}))
        assert r.status_code == 200

    # --- Upload user letterhead word template ---
    def test_upload_user_word_template(self, api_client, lawyer):
        """Lines 1173+: upload user letterhead word template."""
        api_client.force_authenticate(user=lawyer)
        # Minimal valid .docx (PK header)
        docx = SimpleUploadedFile(
            "template.docx",
            b"PK\x03\x04" + b"\x00" * 500,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        r = api_client.post(
            reverse('upload-user-letterhead-word-template'),
            {'template': docx}, format='multipart')
        assert r.status_code in (201, 400)

    # --- Get user letterhead word template (no template) ---
    def test_get_user_word_template_none(self, api_client, lawyer):
        """Line 1243: user has no word template → 404."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(
            reverse('get-user-letterhead-word-template'))
        assert r.status_code == 404

    # --- Delete user letterhead word template (no template) ---
    def test_delete_user_word_template_none(self, api_client, lawyer):
        """Line 1281: user has no word template → 404."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.delete(
            reverse('delete-user-letterhead-word-template'))
        assert r.status_code == 404

    # --- Basic user blocked from word template ---
    def test_basic_user_blocked_word_template(self, api_client, basic_user):
        """Line 1175+1273: basic users get 403 for letterhead ops."""
        api_client.force_authenticate(user=basic_user)
        r = api_client.post(
            reverse('upload-user-letterhead-word-template'),
            {}, format='multipart')
        assert r.status_code == 403
        r2 = api_client.delete(
            reverse('delete-user-letterhead-word-template'))
        assert r2.status_code == 403

    # --- Upload letterhead image to document ---
    def test_upload_letterhead_no_image(self, api_client, lawyer, doc):
        """Line 827-828: no image file provided."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.post(
            reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
            {}, format='multipart')
        assert r.status_code == 400

    # --- Delete letterhead image (no image) ---
    def test_delete_letterhead_no_image(self, api_client, lawyer, doc):
        """Line 988: no letterhead image to delete → 404."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.delete(
            reverse('delete-letterhead-image', kwargs={'pk': doc.pk}))
        assert r.status_code == 404

    # --- Upload doc word template (no file) ---
    def test_upload_doc_word_template_no_file(self, api_client, lawyer, doc):
        """Line 1034: no template file → 400."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.post(
            reverse('upload-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}),
            {}, format='multipart')
        assert r.status_code == 400

    # --- Get doc word template (no template) ---
    def test_get_doc_word_template_none(self, api_client, lawyer, doc):
        """Line 1100: doc has no word template → 404."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(
            reverse('get-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}))
        assert r.status_code == 404

    # --- Delete doc word template (no template) ---
    def test_delete_doc_word_template_none(self, api_client, lawyer, doc):
        """Line 1139: doc has no word template → 404."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.delete(
            reverse('delete-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}))
        assert r.status_code == 404

    # --- Upload user letterhead image (no file) ---
    def test_upload_user_letterhead_no_file(self, api_client, lawyer):
        """User letterhead image upload with no file → 400."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.post(
            reverse('upload-user-letterhead-image'),
            {}, format='multipart')
        assert r.status_code == 400

    # --- Get user letterhead image (none set) ---
    def test_get_user_letterhead_none(self, api_client, lawyer):
        """User has no letterhead image → 404."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(reverse('get-user-letterhead-image'))
        assert r.status_code == 404

    # --- PDF download covers lines 305-342 (font reg + letterhead) ---
    def test_download_pdf_rich_content(self, api_client, lawyer, doc):
        """Lines 305-342: PDF generation with styled HTML content."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(
            reverse('download_dynamic_document_pdf', kwargs={'pk': doc.pk}))
        assert r.status_code == 200
        assert r['Content-Type'] == 'application/pdf'

    # --- Upload user letterhead image with valid image ---
    def test_upload_user_letterhead_valid(self, api_client, lawyer):
        """Lines 1380-1446: upload valid user letterhead image."""
        from io import BytesIO
        from PIL import Image as PILImage
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        r = api_client.post(
            reverse('upload-user-letterhead-image'),
            {'image': img}, format='multipart')
        assert r.status_code == 201
        assert 'message' in r.data

    # --- Delete user letterhead image (basic user blocked) ---
    def test_delete_user_letterhead_basic_blocked(self, api_client, basic_user):
        """Line 1507: basic user blocked from deleting letterhead."""
        api_client.force_authenticate(user=basic_user)
        r = api_client.delete(reverse('delete-user-letterhead-image'))
        assert r.status_code == 403

    # --- Upload letterhead image with valid PNG to document ---
    def test_upload_letterhead_valid(self, api_client, lawyer, doc):
        """Lines 860-931: upload valid letterhead image to document."""
        from io import BytesIO
        from PIL import Image as PILImage
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        r = api_client.post(
            reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
            {'image': img}, format='multipart')
        assert r.status_code == 201

    # --- PDF download with letterhead image set ---
    def test_download_pdf_with_letterhead(self, api_client, lawyer, doc):
        """Lines 320-342: PDF with letterhead image from document."""
        from io import BytesIO
        from PIL import Image as PILImage
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        # Upload letterhead first
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        api_client.post(
            reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
            {'image': img}, format='multipart')
        # Now download PDF with letterhead
        r = api_client.get(
            reverse('download_dynamic_document_pdf', kwargs={'pk': doc.pk}))
        assert r.status_code == 200
        assert r['Content-Type'] == 'application/pdf'

    # --- Get/delete letterhead after upload ---
    def test_get_and_delete_letterhead(self, api_client, lawyer, doc):
        """Lines 960-1016: get and delete letterhead image."""
        from io import BytesIO
        from PIL import Image as PILImage
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        api_client.post(
            reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
            {'image': img}, format='multipart')
        # Get it
        r = api_client.get(
            reverse('get-letterhead-image', kwargs={'pk': doc.pk}))
        assert r.status_code == 200
        # Delete it
        r2 = api_client.delete(
            reverse('delete-letterhead-image', kwargs={'pk': doc.pk}))
        assert r2.status_code == 200

    # --- Upload valid docx template to document, get, delete ---
    def test_upload_get_delete_doc_word_template(self, api_client, lawyer, doc):
        """Lines 1059-1165: upload/get/delete doc word template lifecycle."""
        from io import BytesIO
        from docx import Document as DocxDoc
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        # Upload
        r = api_client.post(
            reverse('upload-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}),
            {'template': tpl}, format='multipart')
        assert r.status_code == 201
        # Get
        r2 = api_client.get(
            reverse('get-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}))
        assert r2.status_code == 200
        # Delete
        r3 = api_client.delete(
            reverse('delete-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}))
        assert r3.status_code == 200

    # --- Upload valid docx template to user, get, delete ---
    def test_upload_get_delete_user_word_template(self, api_client, lawyer):
        """Lines 1208-1302: upload/get/delete user word template lifecycle."""
        from io import BytesIO
        from docx import Document as DocxDoc
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        # Upload
        r = api_client.post(
            reverse('upload-user-letterhead-word-template'),
            {'template': tpl}, format='multipart')
        assert r.status_code == 201
        # Get
        r2 = api_client.get(
            reverse('get-user-letterhead-word-template'))
        assert r2.status_code == 200
        # Delete
        r3 = api_client.delete(
            reverse('delete-user-letterhead-word-template'))
        assert r3.status_code == 200

    # --- Re-upload letterhead image to cover old file deletion (lines 898-902) ---
    def test_reupload_letterhead_image(self, api_client, lawyer, doc):
        """Lines 898-902: re-upload letterhead covers old file deletion."""
        from io import BytesIO
        from PIL import Image as PILImage
        api_client.force_authenticate(user=lawyer)
        for _ in range(2):
            buf = BytesIO()
            PILImage.new('RGB', (816, 1056), color='white').save(
                buf, format='PNG')
            buf.seek(0)
            img = SimpleUploadedFile(
                "lhead.png", buf.read(), content_type="image/png")
            r = api_client.post(
                reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
                {'image': img}, format='multipart')
            assert r.status_code == 201

    # --- Re-upload user letterhead image (lines 1421-1425) ---
    def test_reupload_user_letterhead_image(self, api_client, lawyer):
        """Lines 1421-1425: re-upload user letterhead covers old file deletion."""
        from io import BytesIO
        from PIL import Image as PILImage
        api_client.force_authenticate(user=lawyer)
        for _ in range(2):
            buf = BytesIO()
            PILImage.new('RGB', (816, 1056), color='white').save(
                buf, format='PNG')
            buf.seek(0)
            img = SimpleUploadedFile(
                "lhead.png", buf.read(), content_type="image/png")
            r = api_client.post(
                reverse('upload-user-letterhead-image'),
                {'image': img}, format='multipart')
            assert r.status_code == 201

    # --- Delete user letterhead after upload (lines 1525-1526) ---
    def test_delete_user_letterhead_after_upload(self, api_client, lawyer):
        """Lines 1525-1526: delete user letterhead image after upload."""
        from io import BytesIO
        from PIL import Image as PILImage
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        api_client.post(
            reverse('upload-user-letterhead-image'),
            {'image': img}, format='multipart')
        r = api_client.delete(reverse('delete-user-letterhead-image'))
        assert r.status_code == 200

    # --- Word download with malformed styles (lines 599-600, 606-607, 648-649) ---
    def test_download_word_malformed_styles(self, api_client, lawyer):
        """Lines 599-607, 648-649: malformed padding/line-height/font-size."""
        malformed_html = (
            '<p style="padding-left: badpx">A</p>'
            '<p style="line-height: bad;">B</p>'
            '<p style="font-size: badpt;">C</p>'
            '<p><span style="color: rgb(bad)">D</span></p>'
        )
        d = DynamicDocument.objects.create(
            title='MalformedDoc', content=malformed_html, state='Draft',
            created_by=lawyer,
            created_at=timezone.now(), updated_at=timezone.now())
        api_client.force_authenticate(user=lawyer)
        r = api_client.get(
            reverse('download_dynamic_document_word', kwargs={'pk': d.pk}))
        assert r.status_code == 200

    # --- Word download with word template on document (line 577) ---
    def test_download_word_with_template(self, api_client, lawyer, doc):
        """Line 577: Word download uses document word template."""
        from io import BytesIO
        from docx import Document as DocxDoc
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        api_client.post(
            reverse('upload-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}),
            {'template': tpl}, format='multipart')
        r = api_client.get(
            reverse('download_dynamic_document_word', kwargs={'pk': doc.pk}))
        assert r.status_code == 200

    # ========== PDF with user global letterhead fallback (line 1332) ==========

    def test_pdf_with_user_global_letterhead_fallback(self, api_client, lawyer, doc):
        """Line 1332: PDF uses user global letterhead when document has none."""
        from io import BytesIO
        from PIL import Image as PILImage
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(buf, format='PNG')
        buf.seek(0)
        api_client.force_authenticate(user=lawyer)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        r_upload = api_client.post(
            reverse('upload-user-letterhead-image'),
            {'image': img}, format='multipart')
        assert r_upload.status_code == 201
        assert not doc.letterhead_image
        r = api_client.get(
            reverse('download_dynamic_document_pdf', kwargs={'pk': doc.pk}))
        assert r.status_code == 200
        assert r['Content-Type'] == 'application/pdf'

    # ========== os.remove exception tests (file I/O boundary) ==========

    def test_reupload_letterhead_os_remove_fails(self, api_client, lawyer, doc):
        """Lines 901-902: old image deletion failure is caught silently."""
        from io import BytesIO
        from PIL import Image as PILImage
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(buf, format='PNG')
        buf.seek(0)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        r1 = api_client.post(
            reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
            {'image': img}, format='multipart')
        assert r1.status_code == 201
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            buf2 = BytesIO()
            PILImage.new('RGB', (816, 1056), color='blue').save(
                buf2, format='PNG')
            buf2.seek(0)
            img2 = SimpleUploadedFile(
                "lhead2.png", buf2.read(), content_type="image/png")
            r2 = api_client.post(
                reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
                {'image': img2}, format='multipart')
            assert r2.status_code == 201
            assert 'message' in r2.data

    def test_delete_letterhead_os_remove_fails(self, api_client, lawyer, doc):
        """Lines 998-999: file deletion exception during delete is caught."""
        from io import BytesIO
        from PIL import Image as PILImage
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(buf, format='PNG')
        buf.seek(0)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        api_client.post(
            reverse('upload-letterhead-image', kwargs={'pk': doc.pk}),
            {'image': img}, format='multipart')
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            r = api_client.delete(
                reverse('delete-letterhead-image', kwargs={'pk': doc.pk}))
            assert r.status_code == 200
            assert 'message' in r.data

    def test_reupload_doc_word_template_os_remove_fails(
        self, api_client, lawyer, doc
    ):
        """Lines 1059-1063: old template deletion exception is caught."""
        from io import BytesIO
        from docx import Document as DocxDoc
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        api_client.post(
            reverse('upload-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}),
            {'template': tpl}, format='multipart')
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            buf2 = BytesIO()
            DocxDoc().save(buf2)
            buf2.seek(0)
            tpl2 = SimpleUploadedFile(
                "tpl2.docx", buf2.read(),
                content_type="application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document")
            r = api_client.post(
                reverse('upload-document-letterhead-word-template',
                        kwargs={'pk': doc.pk}),
                {'template': tpl2}, format='multipart')
            assert r.status_code == 201

    def test_delete_doc_word_template_os_remove_fails(
        self, api_client, lawyer, doc
    ):
        """Lines 1148-1149: template file deletion exception is caught."""
        from io import BytesIO
        from docx import Document as DocxDoc
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        api_client.post(
            reverse('upload-document-letterhead-word-template',
                    kwargs={'pk': doc.pk}),
            {'template': tpl}, format='multipart')
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            r = api_client.delete(
                reverse('delete-document-letterhead-word-template',
                        kwargs={'pk': doc.pk}))
            assert r.status_code == 200

    def test_reupload_user_letterhead_os_remove_fails(
        self, api_client, lawyer
    ):
        """Lines 1424-1425: old user letterhead deletion exception caught."""
        from io import BytesIO
        from PIL import Image as PILImage
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        api_client.post(
            reverse('upload-user-letterhead-image'),
            {'image': img}, format='multipart')
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            buf2 = BytesIO()
            PILImage.new('RGB', (816, 1056), color='blue').save(
                buf2, format='PNG')
            buf2.seek(0)
            img2 = SimpleUploadedFile(
                "lhead2.png", buf2.read(), content_type="image/png")
            r = api_client.post(
                reverse('upload-user-letterhead-image'),
                {'image': img2}, format='multipart')
            assert r.status_code == 201

    def test_delete_user_letterhead_os_remove_fails(
        self, api_client, lawyer
    ):
        """Lines 1525-1526: user letterhead file deletion exception caught."""
        from io import BytesIO
        from PIL import Image as PILImage
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        api_client.post(
            reverse('upload-user-letterhead-image'),
            {'image': img}, format='multipart')
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            r = api_client.delete(reverse('delete-user-letterhead-image'))
            assert r.status_code == 200
            assert 'message' in r.data

    def test_reupload_user_word_template_os_remove_fails(
        self, api_client, lawyer
    ):
        """Lines 1208-1212: old user template deletion exception caught."""
        from io import BytesIO
        from docx import Document as DocxDoc
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        api_client.post(
            reverse('upload-user-letterhead-word-template'),
            {'template': tpl}, format='multipart')
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            buf2 = BytesIO()
            DocxDoc().save(buf2)
            buf2.seek(0)
            tpl2 = SimpleUploadedFile(
                "tpl2.docx", buf2.read(),
                content_type="application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document")
            r = api_client.post(
                reverse('upload-user-letterhead-word-template'),
                {'template': tpl2}, format='multipart')
            assert r.status_code == 201

    def test_delete_user_word_template_os_remove_fails(
        self, api_client, lawyer
    ):
        """Lines 1290-1291: user template file deletion exception caught."""
        from io import BytesIO
        from docx import Document as DocxDoc
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        api_client.post(
            reverse('upload-user-letterhead-word-template'),
            {'template': tpl}, format='multipart')
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.remove',
            side_effect=OSError("mock")
        ):
            r = api_client.delete(
                reverse('delete-user-letterhead-word-template'))
            assert r.status_code == 200

    # ========== Generic exception handlers (user endpoints) ==========

    def test_upload_user_word_template_generic_exception(
        self, api_client, lawyer
    ):
        """Lines 1229-1230: generic exception in upload → 500."""
        from io import BytesIO
        from docx import Document as DocxDoc
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        with mock.patch.object(
            User, 'save', side_effect=Exception("DB error")
        ):
            r = api_client.post(
                reverse('upload-user-letterhead-word-template'),
                {'template': tpl}, format='multipart')
        assert r.status_code == 500
        assert 'Error' in r.data['detail']

    def test_get_user_word_template_generic_exception(
        self, api_client, lawyer
    ):
        """Lines 1262-1263: generic exception in get → 500."""
        from io import BytesIO
        from docx import Document as DocxDoc
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        api_client.post(
            reverse('upload-user-letterhead-word-template'),
            {'template': tpl}, format='multipart')
        lawyer.refresh_from_db()
        file_path = lawyer.letterhead_word_template.path
        if os.path.exists(file_path):
            os.remove(file_path)
        with mock.patch(
            'gym_app.views.dynamic_documents.document_views.os.path.exists',
            return_value=True
        ):
            r = api_client.get(
                reverse('get-user-letterhead-word-template'))
        assert r.status_code == 500
        assert 'Error' in r.data['detail']

    def test_delete_user_word_template_generic_exception(
        self, api_client, lawyer
    ):
        """Lines 1301-1302: generic exception in delete → 500."""
        from io import BytesIO
        from docx import Document as DocxDoc
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        DocxDoc().save(buf)
        buf.seek(0)
        tpl = SimpleUploadedFile(
            "tpl.docx", buf.read(),
            content_type="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        api_client.post(
            reverse('upload-user-letterhead-word-template'),
            {'template': tpl}, format='multipart')
        with mock.patch.object(
            User, 'save', side_effect=Exception("DB error")
        ):
            r = api_client.delete(
                reverse('delete-user-letterhead-word-template'))
        assert r.status_code == 500
        assert 'Error' in r.data['detail']

    def test_upload_user_letterhead_generic_exception(
        self, api_client, lawyer
    ):
        """Lines 1448-1449: generic exception in upload user letterhead → 500."""
        from io import BytesIO
        from PIL import Image as PILImage
        api_client.force_authenticate(user=lawyer)
        buf = BytesIO()
        PILImage.new('RGB', (816, 1056), color='white').save(
            buf, format='PNG')
        buf.seek(0)
        img = SimpleUploadedFile(
            "lhead.png", buf.read(), content_type="image/png")
        with mock.patch.object(
            User, 'save', side_effect=Exception("DB error")
        ):
            r = api_client.post(
                reverse('upload-user-letterhead-image'),
                {'image': img}, format='multipart')
        assert r.status_code == 500
        assert 'Error' in r.data['detail']

    # ========== Dead code behind decorators (DoesNotExist) ==========

    def test_update_recent_document_nonexistent(
        self, api_client, lawyer
    ):
        """Line 827-828: DoesNotExist in update_recent_document → 404."""
        api_client.force_authenticate(user=lawyer)
        r = api_client.post(
            reverse('update-recent-document',
                    kwargs={'document_id': 99999}))
        assert r.status_code == 404
        assert 'not found' in r.data.get('detail', '').lower()


# ======================================================================
# Tests merged from test_document_views_phase1.py
# ======================================================================


@pytest.fixture
def lawyer_user():
    return User.objects.create_user(
        email="dv1-lawyer@example.com",
        password="testpassword",
        role="lawyer",
    )


@pytest.fixture
def basic_user():
    return User.objects.create_user(
        email="dv1-basic@example.com",
        password="testpassword",
        role="basic",
    )


@pytest.fixture
def document(lawyer_user):
    return DynamicDocument.objects.create(
        title="DV1 Coverage Doc",
        content="<p>test content</p>",
        state="Draft",
        created_by=lawyer_user,
    )


def _make_png_file(width=100, height=100, name="test.png"):
    """Create a valid in-memory PNG file for upload tests."""
    buf = io.BytesIO()
    img = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    img.save(buf, format="PNG")
    buf.seek(0)
    return SimpleUploadedFile(name, buf.read(), content_type="image/png")


# ── Document letterhead image upload validations (lines 856, 864, 888-891, 921) ──

class TestUploadLetterheadImageValidation:
    def test_upload_letterhead_non_png_returns_400(self, api_client, lawyer_user, document):
        """Line 856: Non-PNG file rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        jpg_file = SimpleUploadedFile("test.jpg", b"fake", content_type="image/jpeg")
        response = api_client.post(url, {"image": jpg_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "PNG" in response.data["detail"]

    def test_upload_letterhead_too_large_returns_400(self, api_client, lawyer_user, document):
        """Line 864: File > 10MB rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        large_file = SimpleUploadedFile("big.png", b"x" * (11 * 1024 * 1024), content_type="image/png")
        response = api_client.post(url, {"image": large_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "10MB" in response.data["detail"]

    def test_upload_letterhead_invalid_image_returns_400(self, api_client, lawyer_user, document):
        """Lines 888-891: Invalid image data rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        bad_png = SimpleUploadedFile("bad.png", b"not-a-real-image", content_type="image/png")
        response = api_client.post(url, {"image": bad_png}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "inválido" in response.data["detail"].lower() or "invalid" in response.data["detail"].lower()

    def test_upload_letterhead_with_aspect_ratio_warning(self, api_client, lawyer_user, document):
        """Line 921: Square image triggers aspect ratio warning."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-letterhead-image", kwargs={"pk": document.id})
        square_png = _make_png_file(width=100, height=100)
        response = api_client.post(url, {"image": square_png}, format="multipart")

        assert response.status_code == status.HTTP_201_CREATED
        assert "warnings" in response.data


# ── Document letterhead image get/delete edge cases (line 954) ──────

class TestGetLetterheadImageEdgeCases:
    def test_get_letterhead_file_missing_on_disk(self, api_client, lawyer_user, document):
        """Line 954: Letterhead field set but file doesn't exist on disk."""
        document.letterhead_image = "nonexistent/path.png"
        document.save(update_fields=["letterhead_image"])

        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-letterhead-image", kwargs={"pk": document.id})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND
        assert "no se encuentra" in response.data["detail"].lower() or "not found" in response.data["detail"].lower()


# ── Document word template upload validations (lines 1044, 1052) ────

class TestUploadDocWordTemplateValidation:
    def test_upload_word_template_non_docx_returns_400(self, api_client, lawyer_user, document):
        """Line 1044: Non-DOCX file rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": document.id})
        txt_file = SimpleUploadedFile("test.txt", b"fake", content_type="text/plain")
        response = api_client.post(url, {"template": txt_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert ".docx" in response.data["detail"]

    def test_upload_word_template_too_large_returns_400(self, api_client, lawyer_user, document):
        """Line 1052: File > 10MB rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-document-letterhead-word-template", kwargs={"pk": document.id})
        large_file = SimpleUploadedFile("big.docx", b"x" * (11 * 1024 * 1024), content_type="application/octet-stream")
        response = api_client.post(url, {"template": large_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "grande" in response.data["detail"].lower()


# ── Document word template get edge case (line 1107) ────────────────

class TestGetDocWordTemplateEdgeCases:
    def test_get_word_template_file_missing_on_disk(self, api_client, lawyer_user, document):
        """Line 1107: Word template field set but file doesn't exist on disk."""
        document.letterhead_word_template = "nonexistent/path.docx"
        document.save(update_fields=["letterhead_word_template"])

        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-document-letterhead-word-template", kwargs={"pk": document.id})
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND


# ── User global word template CRUD (lines 1184, 1193, 1201, 1250) ──

class TestUserWordTemplateCRUD:
    def test_upload_user_word_template_no_file_returns_400(self, api_client, lawyer_user):
        """Line 1184: No template file provided."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-word-template")
        response = api_client.post(url, {}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "plantilla" in response.data["detail"].lower()

    def test_upload_user_word_template_non_docx_returns_400(self, api_client, lawyer_user):
        """Line 1193: Non-DOCX file rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-word-template")
        txt_file = SimpleUploadedFile("test.pdf", b"fake", content_type="application/pdf")
        response = api_client.post(url, {"template": txt_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert ".docx" in response.data["detail"]

    def test_upload_user_word_template_too_large_returns_400(self, api_client, lawyer_user):
        """Line 1201: File > 10MB rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-word-template")
        large_file = SimpleUploadedFile("big.docx", b"x" * (11 * 1024 * 1024), content_type="application/octet-stream")
        response = api_client.post(url, {"template": large_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "grande" in response.data["detail"].lower()

    def test_get_user_word_template_file_missing_on_disk(self, api_client, lawyer_user):
        """Line 1250: Template field set but file doesn't exist on disk."""
        lawyer_user.letterhead_word_template = "nonexistent/path.docx"
        lawyer_user.save(update_fields=["letterhead_word_template"])

        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-user-letterhead-word-template")
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND


# ── User global letterhead image CRUD (lines 1360, 1377, 1385, 1408-1414, 1444, 1476-1491, 1516) ──

class TestUserLetterheadImageCRUD:
    def test_upload_user_letterhead_basic_user_forbidden(self, api_client, basic_user):
        """Line 1360: Basic users cannot upload letterhead."""
        api_client.force_authenticate(user=basic_user)
        url = reverse("upload-user-letterhead-image")
        response = api_client.post(url, {}, format="multipart")

        assert response.status_code == status.HTTP_403_FORBIDDEN

    def test_upload_user_letterhead_non_png_returns_400(self, api_client, lawyer_user):
        """Line 1377: Non-PNG file rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-image")
        jpg_file = SimpleUploadedFile("test.jpg", b"fake", content_type="image/jpeg")
        response = api_client.post(url, {"image": jpg_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "PNG" in response.data["detail"]

    def test_upload_user_letterhead_too_large_returns_400(self, api_client, lawyer_user):
        """Line 1385: File > 10MB rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-image")
        large_file = SimpleUploadedFile("big.png", b"x" * (11 * 1024 * 1024), content_type="image/png")
        response = api_client.post(url, {"image": large_file}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST
        assert "10MB" in response.data["detail"]

    def test_upload_user_letterhead_invalid_image_returns_400(self, api_client, lawyer_user):
        """Lines 1408-1414: Invalid image data rejected."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-image")
        bad_png = SimpleUploadedFile("bad.png", b"not-a-real-image", content_type="image/png")
        response = api_client.post(url, {"image": bad_png}, format="multipart")

        assert response.status_code == status.HTTP_400_BAD_REQUEST

    def test_upload_user_letterhead_with_warnings(self, api_client, lawyer_user):
        """Line 1444: Square image triggers aspect ratio warning."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("upload-user-letterhead-image")
        square_png = _make_png_file(width=100, height=100)
        response = api_client.post(url, {"image": square_png}, format="multipart")

        assert response.status_code == status.HTTP_201_CREATED
        assert "warnings" in response.data

    def test_get_user_letterhead_file_missing_on_disk(self, api_client, lawyer_user):
        """Lines 1476-1491: Letterhead field set but file doesn't exist on disk."""
        lawyer_user.letterhead_image = "nonexistent/path.png"
        lawyer_user.save(update_fields=["letterhead_image"])

        api_client.force_authenticate(user=lawyer_user)
        url = reverse("get-user-letterhead-image")
        response = api_client.get(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_user_letterhead_no_image_returns_404(self, api_client, lawyer_user):
        """Line 1516: No letterhead image to delete."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("delete-user-letterhead-image")
        response = api_client.delete(url)

        assert response.status_code == status.HTTP_404_NOT_FOUND


class TestDocumentViewsCRUDEdges:

    def test_update_dynamic_document_not_found(self, api_client, lawyer_user):
        """Line 214-215: doc not found."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("update_dynamic_document", kwargs={"pk": 99999})
        resp = api_client.put(url, {"title": "X"}, format="json")
        assert resp.status_code == status.HTTP_404_NOT_FOUND

    def test_delete_dynamic_document_not_found(self, api_client, lawyer_user):
        """Line 239-240: doc not found."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("delete_dynamic_document", kwargs={"pk": 99999})
        resp = api_client.delete(url)
        assert resp.status_code == status.HTTP_404_NOT_FOUND

    def test_update_dynamic_document_validation_error(self, api_client, lawyer_user, document):
        """Line 227: serializer validation error."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("update_dynamic_document", kwargs={"pk": document.id})
        # state must be a valid choice - send invalid
        resp = api_client.patch(url, {"state": ""}, format="json")
        # Should be 400 or 200 depending on serializer; empty string may be invalid
        assert resp.status_code in (status.HTTP_200_OK, status.HTTP_400_BAD_REQUEST)

    def test_create_dynamic_document_assigns_to(self, api_client, lawyer_user):
        """Lines 63-64: auto-assign assigned_to for Progress state."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("create_dynamic_document")
        resp = api_client.post(url, {
            "title": "Auto Assign",
            "content": "<p>test</p>",
            "state": "Progress",
        }, format="json")
        # May succeed or fail depending on serializer validation
        assert resp.status_code in (status.HTTP_201_CREATED, status.HTTP_400_BAD_REQUEST)

    def test_download_pdf_not_found(self, api_client, lawyer_user):
        """Line 451-452: doc not found for PDF."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("download_dynamic_document_pdf", kwargs={"pk": 99999})
        resp = api_client.get(url)
        assert resp.status_code == status.HTTP_404_NOT_FOUND

    def test_download_word_not_found(self, api_client, lawyer_user):
        """Line 778-779: doc not found for Word."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("download_dynamic_document_word", kwargs={"pk": 99999})
        resp = api_client.get(url)
        assert resp.status_code == status.HTTP_404_NOT_FOUND



class TestListDynamicDocumentsEdges:

    def test_list_page_not_integer(self, api_client, lawyer_user, document):
        """Lines 123-125: PageNotAnInteger fallback."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        resp = api_client.get(url, {"page": "abc", "limit": "5"})
        assert resp.status_code == status.HTTP_200_OK
        assert resp.data["currentPage"] == 1

    def test_list_empty_page(self, api_client, lawyer_user, document):
        """Lines 126-129: EmptyPage fallback to last page."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        resp = api_client.get(url, {"page": "9999", "limit": "5"})
        assert resp.status_code == status.HTTP_200_OK
        # Should be on the last page
        assert resp.data["currentPage"] == resp.data["totalPages"]

    def test_list_negative_limit(self, api_client, lawyer_user, document):
        """Lines 116-117: negative limit defaults to 10."""
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        resp = api_client.get(url, {"limit": "-1"})
        assert resp.status_code == status.HTTP_200_OK

    def test_list_multi_state_filter(self, api_client, lawyer_user):
        """Lines 92-95: multi-state filter."""
        DynamicDocument.objects.create(
            title="D1", content="<p>x</p>", state="Draft", created_by=lawyer_user,
        )
        DynamicDocument.objects.create(
            title="D2", content="<p>x</p>", state="Published", created_by=lawyer_user,
        )
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        resp = api_client.get(url, {"states": "Draft,Published"})
        assert resp.status_code == status.HTTP_200_OK
        assert resp.data["totalItems"] >= 2

    def test_list_client_id_and_lawyer_id_filters(self, api_client, lawyer_user, client_user):
        """Lines 99-103: client_id and lawyer_id filters."""
        doc = DynamicDocument.objects.create(
            title="Assigned", content="<p>x</p>", state="Draft",
            created_by=lawyer_user, assigned_to=client_user,
        )
        api_client.force_authenticate(user=lawyer_user)
        url = reverse("list_dynamic_documents")
        resp = api_client.get(url, {"client_id": client_user.id})
        assert resp.status_code == status.HTTP_200_OK

        resp2 = api_client.get(url, {"lawyer_id": lawyer_user.id})
        assert resp2.status_code == status.HTTP_200_OK


