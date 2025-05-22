"""
Utility functions for handling document signatures.

This module contains functions to embed signatures in documents when 
generating PDF or Word files, as well as helper functions for validating
and processing signature-related data.
"""
import os
import datetime
from django.conf import settings
from reportlab.lib.units import inch
from reportlab.platypus import Image
from docx.shared import Inches
from PIL import Image as PILImage
from io import BytesIO


def get_signature_image_path(user):
    """
    Get the file path of a user's signature image.
    
    Args:
        user: User object containing signature relation
        
    Returns:
        str: Path to the signature image or None if not available
    """
    if not hasattr(user, 'signature') or not user.signature:
        return None
        
    signature_path = user.signature.signature_image.path
    if not os.path.exists(signature_path):
        return None
        
    return signature_path


def add_signatures_to_pdf_context(document, html_content):
    """
    Add signature placeholders to the PDF HTML content.
    
    Args:
        document: DynamicDocument object
        html_content: HTML content for the PDF
        
    Returns:
        str: HTML content with signature placeholders
    """
    if not document.requires_signature:
        return html_content
        
    # Get signed signatures
    signatures = document.signatures.filter(signed=True).order_by('signature_position')
    if not signatures:
        return html_content
        
    # Create signature HTML section
    signature_html = """
    <div style="margin-top: 40px; border-top: 1px solid #ccc; padding-top: 20px;">
        <h3 style="font-size: 14pt; margin-bottom: 20px;">Signatures:</h3>
        <div style="display: flex; flex-wrap: wrap; justify-content: space-between;">
    """
    
    for signature in signatures:
        signature_path = get_signature_image_path(signature.signer)
        if not signature_path:
            continue
            
        signer_name = f"{signature.signer.first_name} {signature.signer.last_name}".strip()
        if not signer_name:
            signer_name = signature.signer.email
            
        signed_date = signature.signed_at.strftime("%d/%m/%Y %H:%M") if signature.signed_at else ""
        
        signature_html += f"""
        <div style="width: 45%; margin-bottom: 20px;">
            <img src="{signature_path}" width="200" height="80" />
            <div style="margin-top: 5px; font-size: 11pt;">{signer_name}</div>
            <div style="font-size: 9pt; color: #666;">Signed: {signed_date}</div>
        </div>
        """
    
    signature_html += """
        </div>
    </div>
    """
    
    # Append signature section to the end of the body
    modified_html = html_content.replace("</body>", f"{signature_html}</body>")
    return modified_html


def add_signatures_to_word_document(document, doc_obj):
    """
    Add signatures to a Word document.
    
    Args:
        document: DynamicDocument object
        doc_obj: python-docx Document object
        
    Returns:
        Document: Updated Word document with signatures
    """
    if not document.requires_signature:
        return doc_obj
        
    # Get signed signatures
    signatures = document.signatures.filter(signed=True).order_by('signature_position')
    if not signatures:
        return doc_obj
        
    # Add a separator
    doc_obj.add_paragraph()
    doc_obj.add_paragraph("_" * 71)
    doc_obj.add_paragraph()
    
    # Add signature heading
    sig_heading = doc_obj.add_heading("Signatures:", level=3)
    
    # Create a table for signatures (2 columns)
    sig_count = signatures.count()
    rows_needed = (sig_count + 1) // 2  # Calculate rows needed (rounding up)
    table = doc_obj.add_table(rows=rows_needed, cols=2)
    
    # Add signatures to the table
    row_idx = 0
    col_idx = 0
    
    for signature in signatures:
        signature_path = get_signature_image_path(signature.signer)
        if not signature_path:
            continue
            
        # Get cell and add signature image
        cell = table.cell(row_idx, col_idx)
        
        # Add signature image
        try:
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(signature_path, width=Inches(2))
            
            # Add signer name
            signer_name = f"{signature.signer.first_name} {signature.signer.last_name}".strip()
            if not signer_name:
                signer_name = signature.signer.email
            
            name_paragraph = cell.add_paragraph(signer_name)
            
            # Add signature date
            signed_date = signature.signed_at.strftime("%d/%m/%Y %H:%M") if signature.signed_at else ""
            date_paragraph = cell.add_paragraph(f"Signed: {signed_date}")
            
            # Move to next position
            col_idx = (col_idx + 1) % 2
            if col_idx == 0:
                row_idx += 1
                
        except Exception as e:
            print(f"Error adding signature to Word document: {str(e)}")
            continue
    
    return doc_obj 