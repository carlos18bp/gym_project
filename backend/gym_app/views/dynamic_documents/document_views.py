import io
import re
import os
from django.conf import settings
from django.http import FileResponse, Http404
from django.template.loader import get_template
from PIL import Image
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from bs4 import BeautifulSoup, NavigableString
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from gym_app.models.dynamic_document import DynamicDocument, RecentDocument
from gym_app.serializers.dynamic_document import DynamicDocumentSerializer, RecentDocumentSerializer
from django.utils import timezone
from .permissions import (
    require_document_visibility,
    require_document_visibility_by_id,
    require_document_usability,
    filter_documents_by_visibility
)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_dynamic_document(request):
    """
    Create a new dynamic document with optional permissions.
    
    You can now create documents and set permissions in a single API call:
    
    Basic fields:
    - title, content, state, is_public, requires_signature
    - variables: array of document variables
    - tags: array of tag IDs
    - signers: array of user IDs for signatures
    
    Permission fields (optional):
    - visibility_user_ids: array of user IDs who can view the document
    - usability_user_ids: array of user IDs who can use/edit the document
    
    Example payload:
    {
        "title": "Contract Template",
        "content": "Document content...",
        "is_public": false,
        "visibility_user_ids": [5, 6, 7, 8],
        "usability_user_ids": [5, 6],
        "variables": [...],
        "tags": [1, 2]
    }
    """
    # If it's a creation from the client, assign `assigned_to`.
    if request.data.get('state') in ['Progress', 'Completed'] and not request.data.get('assigned_to'):
        request.data['assigned_to'] = request.user.id

    # Validate and save the document
    serializer = DynamicDocumentSerializer(data=request.data, context={'request': request})
    if serializer.is_valid():
        instance = serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@filter_documents_by_visibility
def list_dynamic_documents(request):
    """
    Get a list of all dynamic documents.
    """
    documents = DynamicDocument.objects.prefetch_related('variables', 'tags').all()
    serializer = DynamicDocumentSerializer(documents, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@require_document_visibility
def get_dynamic_document(request, pk):
    """
    Get a specific dynamic document by ID.
    """
    try:
        document = DynamicDocument.objects.prefetch_related(
            'variables',
            'signatures__signer',
            'tags'
        ).get(pk=pk)
        
        # Ensure variables have select_options initialized
        for variable in document.variables.all():
            if variable.field_type == 'select' and not variable.select_options:
                variable.select_options = []
                variable.save()
        
        serializer = DynamicDocumentSerializer(document)
        return Response(serializer.data, status=status.HTTP_200_OK)
    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Dynamic document not found.'}, status=status.HTTP_404_NOT_FOUND)

@api_view(['PUT', 'PATCH'])
@permission_classes([IsAuthenticated])
@require_document_usability('usability')
def update_dynamic_document(request, pk):
    """
    Update an existing dynamic document with optional permissions.
    
    You can now update documents and modify permissions in a single API call:
    
    Basic fields that can be updated:
    - title, content, state, is_public, requires_signature
    - variables: array of document variables (replaces existing)
    - tags: array of tag IDs (replaces existing)
    - signers: array of user IDs for signatures (adds new ones)
    
    Permission fields (optional - replaces existing permissions):
    - visibility_user_ids: array of user IDs who can view the document
    - usability_user_ids: array of user IDs who can use/edit the document
    
    Example payload:
    {
        "title": "Updated Contract Template",
        "is_public": false,
        "visibility_user_ids": [5, 6, 8, 9],
        "usability_user_ids": [5, 6],
        "variables": [...]
    }
    
    Note: Providing permission fields will REPLACE existing permissions.
    """
    try:
        # Get the document and load its related variables
        document = DynamicDocument.objects.prefetch_related('variables', 'tags').get(pk=pk)
    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Dynamic document not found.'}, status=status.HTTP_404_NOT_FOUND)

    # Prevent modifying the `created_by` field
    if 'created_by' in request.data:
        request.data.pop('created_by')

    # Validate and update the document
    serializer = DynamicDocumentSerializer(document, data=request.data, partial=True, context={'request': request})
    if serializer.is_valid():
        instance = serializer.save()
        return Response(serializer.data, status=status.HTTP_200_OK)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
@require_document_usability('owner')
def delete_dynamic_document(request, pk):
    """
    Delete a dynamic document.
    """
    try:
        document = DynamicDocument.objects.get(pk=pk)
    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Dynamic document not found.'}, status=status.HTTP_404_NOT_FOUND)

    document.delete()
    return Response({'detail': 'Dynamic document deleted successfully.'}, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@require_document_visibility
def download_dynamic_document_pdf(request, pk, for_version=False):
    """
    Generates and returns a PDF file for a given document using ReportLab and xhtml2pdf.

    This function retrieves a document from the database, replaces dynamic variables within 
    its content, applies a predefined font style, and converts the content into a properly 
    formatted PDF file. The generated PDF is then returned as a downloadable response.

    Parameters:
        request (HttpRequest): The request object.
        pk (int): The primary key of the document to be retrieved.
        for_version (bool): If True, returns the PDF buffer instead of a FileResponse (for version creation).

    Returns:
        FileResponse: A downloadable PDF file response.
        BytesIO: PDF buffer if for_version is True.
        Response: A JSON response with an error message if an exception occurs.

    Raises:
        FileNotFoundError: If any of the required font files are missing.
        Exception: If there is an error during the HTML-to-PDF conversion or any other 
                   unexpected issue.
    """
    try:
        # Retrieve the document from the database
        document = DynamicDocument.objects.prefetch_related('variables', 'signatures__signer', 'tags').get(pk=pk)

        # Replace variables within the content
        processed_content = document.content
        for variable in document.variables.all():
            processed_content = processed_content.replace(f"{{{{{variable.name_en}}}}}", variable.value or "")

        # Convert HTML to XHTML using BeautifulSoup
        soup = BeautifulSoup(processed_content, 'html.parser')

        # Create the PDF buffer
        pdf_buffer = io.BytesIO()

        # Define font file paths
        font_dir = os.path.abspath(os.path.join(settings.BASE_DIR, 'static', 'fonts'))
        font_paths = {
            "Carlito-Regular": os.path.join(font_dir, "Carlito-Regular.ttf"),
            "Carlito-Bold": os.path.join(font_dir, "Carlito-Bold.ttf"),
            "Carlito-Italic": os.path.join(font_dir, "Carlito-Italic.ttf"),
            "Carlito-BoldItalic": os.path.join(font_dir, "Carlito-BoldItalic.ttf"),
        }

        # Verify that all font files exist
        for name, path in font_paths.items():
            if not os.path.exists(path):
                raise FileNotFoundError(f"Font file not found: {path}")

        # Register fonts in ReportLab
        try:
            pdfmetrics.registerFont(TTFont('Carlito', font_paths["Carlito-Regular"]))
            pdfmetrics.registerFont(TTFont('Carlito-Bold', font_paths["Carlito-Bold"]))
            pdfmetrics.registerFont(TTFont('Carlito-Italic', font_paths["Carlito-Italic"]))
            pdfmetrics.registerFont(TTFont('Carlito-BoldItalic', font_paths["Carlito-BoldItalic"]))
        except Exception as e:
            raise

        # Define background image style if letterhead exists
        background_style = ""
        letterhead_image = get_letterhead_for_document(document, request.user)
        if letterhead_image:
            try:
                # Get the absolute path to the letterhead image
                letterhead_path = os.path.abspath(letterhead_image.path)
                if os.path.exists(letterhead_path):
                    # Convert image to base64 for better xhtml2pdf compatibility
                    import base64
                    with open(letterhead_path, 'rb') as img_file:
                        img_data = base64.b64encode(img_file.read()).decode()
                        img_mime = 'image/png'  # Assuming PNG as per validation
                        background_style = f"""
            background-image: url('data:{img_mime};base64,{img_data}');
            background-repeat: no-repeat;
            background-position: center;
            background-size: contain;
            background-attachment: fixed;"""
            except (ValueError, AttributeError, IOError):
                # Image file doesn't exist or path is invalid
                background_style = ""

        # Define CSS styles for PDF (force Letter size: 8.5 x 11 inches)
        styles = f"""
        <style>
        @page {{
            size: letter;
            margin: 2cm;{background_style}
        }}

        @font-face {{
            font-family: 'Carlito';
            src: url('{font_paths["Carlito-Regular"]}') format('truetype');
            font-weight: normal;
            font-style: normal;
        }}

        @font-face {{
            font-family: 'Carlito';
            src: url('{font_paths["Carlito-Bold"]}') format('truetype');
            font-weight: bold;
            font-style: normal;
        }}

        @font-face {{
            font-family: 'Carlito';
            src: url('{font_paths["Carlito-Italic"]}') format('truetype');
            font-weight: normal;
            font-style: italic;
        }}

        @font-face {{
            font-family: 'Carlito';
            src: url('{font_paths["Carlito-BoldItalic"]}') format('truetype');
            font-weight: bold;
            font-style: italic;
        }}

        body {{
            font-family: 'Carlito', sans-serif !important;
            font-size: 12pt;
        }}

        p, span {{
            font-family: 'Carlito', sans-serif !important;
        }}

        strong {{
            font-weight: bold !important;
            font-family: 'Carlito', sans-serif !important;
        }}

        em {{
            font-style: italic !important;
            font-family: 'Carlito', sans-serif !important;
        }}

        strong em {{
            font-weight: bold !important;
            font-style: italic !important;
            font-family: 'Carlito', sans-serif !important;
        }}

        u {{
            text-decoration: underline !important;
        }}
        </style>
        """

        # Construct the final HTML for the PDF
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{document.title}</title>
            {styles}
        </head>
        <body>
            {soup.prettify()}
        </body>
        </html>
        """

        # Generate the PDF with xhtml2pdf
        pisa_status = pisa.CreatePDF(
            html_content.encode('utf-8'),
            dest=pdf_buffer
        )

        # Check for errors in PDF generation
        if pisa_status.err:
            raise Exception("HTML to PDF conversion failed")

        # Return the generated PDF as a response
        pdf_buffer.seek(0)

        # If this is for a version, return the buffer
        if for_version:
            return pdf_buffer

        return FileResponse(
            pdf_buffer,
            as_attachment=True,
            filename=f"{document.title}.pdf",
            content_type='application/pdf'
        )

    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Document not found.'}, status=status.HTTP_404_NOT_FOUND)
    except FileNotFoundError as e:
        return Response({'detail': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        return Response({'detail': f'Error generating PDF: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
@require_document_visibility
def download_dynamic_document_word(request, pk):
    """
    Generates and returns a Word (.docx) file for the given document using python-docx.
    The document content is retrieved from the database, and its variables are dynamically replaced.
    The content is then parsed from HTML using BeautifulSoup and converted into a Word document,
    applying appropriate formatting, including headings, paragraphs, and styles.
    
    Parameters:
        request (HttpRequest): The HTTP request object.
        pk (int): The primary key of the document to be retrieved.
    
    Returns:
        FileResponse: A response containing the generated Word document.
    """
    try:
        # Retrieve the document from the database
        document = DynamicDocument.objects.prefetch_related('variables', 'signatures__signer', 'tags').get(pk=pk)

        # Replace variables dynamically
        def replace_variables(text):
            processed_text = text
            for variable in document.variables.all():
                pattern = re.compile(rf"{{{{{variable.name_en}}}}}")
                processed_text = pattern.sub(variable.value or "", processed_text)
            return processed_text

        processed_content = replace_variables(document.content)
        
        # Render HTML with template
        template = get_template("pdf_template.html")
        html_content = template.render({"content": processed_content})

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        # Create Word document
        doc = Document()
        
        # Configure page size to Letter (Carta)
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        section = doc.sections[0]
        section.page_width = Inches(8.5)   # 8.5 inches width
        section.page_height = Inches(11)   # 11 inches height (Letter/Carta)
        
        # Configure margins (adjust as needed)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add letterhead image as background using VML in header
        letterhead_image = get_letterhead_for_document(document, request.user)
        if letterhead_image:
            try:
                import os
                import base64
                from PIL import Image as PILImage
                
                # Verify image file exists and is accessible
                letterhead_path = os.path.abspath(letterhead_image.path)
                if not os.path.exists(letterhead_path):
                    raise FileNotFoundError("Letterhead image file not accessible")
                
                # Read and encode image to base64
                with open(letterhead_path, 'rb') as img_file:
                    image_data = img_file.read()
                    image_base64 = base64.b64encode(image_data).decode('utf-8')
                
                # Detect image format
                with PILImage.open(letterhead_path) as img:
                    img_format = img.format.lower() if img.format else 'png'
                
                # Access the header
                header = section.header
                
                # Create or get header paragraph
                if header.paragraphs:
                    header_para = header.paragraphs[0]
                    header_para.clear()
                else:
                    header_para = header.add_paragraph()
                
                # Add the image to document relationships
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                
                # Add image part to document
                image_part = doc.part.package.image_parts.add_image(letterhead_path)
                rId = doc.part.relate_to(image_part, RT.IMAGE)
                
                # Now we need to create a picture element that goes behind text
                # We'll manipulate the XML directly
                from lxml import etree
                from docx.oxml.ns import qn
                
                # Create a run
                run = header_para.add_run()
                
                # Create the picture XML with anchor positioning
                # This is the key: we create a properly formatted anchor element
                pic_xml = f'''
                    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" 
                                   relativeHeight="0" behindDoc="1" locked="0" layoutInCell="0" 
                                   allowOverlap="1" 
                                   xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
                            <wp:simplePos x="0" y="0"/>
                            <wp:positionH relativeFrom="page">
                                <wp:posOffset>0</wp:posOffset>
                            </wp:positionH>
                            <wp:positionV relativeFrom="page">
                                <wp:posOffset>0</wp:posOffset>
                            </wp:positionV>
                            <wp:extent cx="7772400" cy="12700800"/>
                            <wp:effectExtent l="0" t="0" r="0" b="0"/>
                            <wp:wrapNone/>
                            <wp:docPr id="1" name="Background"/>
                            <wp:cNvGraphicFramePr>
                                <a:graphicFrameLocks noChangeAspect="1" 
                                    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"/>
                            </wp:cNvGraphicFramePr>
                            <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                                <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                                    <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
                                        <pic:nvPicPr>
                                            <pic:cNvPr id="1" name="Background"/>
                                            <pic:cNvPicPr/>
                                        </pic:nvPicPr>
                                        <pic:blipFill>
                                            <a:blip r:embed="{rId}" 
                                                xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
                                            <a:stretch>
                                                <a:fillRect/>
                                            </a:stretch>
                                        </pic:blipFill>
                                        <pic:spPr>
                                            <a:xfrm>
                                                <a:off x="0" y="0"/>
                                                <a:ext cx="7772400" cy="12700800"/>
                                            </a:xfrm>
                                            <a:prstGeom prst="rect">
                                                <a:avLst/>
                                            </a:prstGeom>
                                        </pic:spPr>
                                    </pic:pic>
                                </a:graphicData>
                            </a:graphic>
                        </wp:anchor>
                    </w:drawing>
                '''
                
                # Parse the XML string
                drawing_element = etree.fromstring(pic_xml)
                
                # Append to run
                run._element.append(drawing_element)
                
                # Configure header paragraph to take no space
                header_para.paragraph_format.space_after = Pt(0)
                header_para.paragraph_format.space_before = Pt(0)
                header_para.paragraph_format.line_spacing = 1
                
                # Set header distance to 0
                section.header_distance = Inches(0)
                
            except Exception as e:
                # Alternative approach: Add image to first paragraph of body
                try:
                    
                    # Get the first paragraph or create one
                    if not doc.paragraphs:
                        first_para = doc.add_paragraph()
                    else:
                        first_para = doc.paragraphs[0]
                    
                    # Add the image
                    run = first_para.add_run()
                    picture = run.add_picture(letterhead_path, width=Inches(7.5))
                    
                    # Try to access the inline element and modify it
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import qn
                    
                    inline = picture._inline
                    
                    # Get the extent (size) from inline
                    extent = inline.extent
                    cx = extent.cx
                    cy = extent.cy
                    
                    # Get the docPr from inline
                    docPr = inline.docPr
                    
                    # Get graphic from inline
                    graphic = inline.graphic
                    
                    # Create an anchor element
                    anchor_xml = f'''
                        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                                   distT="0" distB="0" distL="0" distR="0" 
                                   simplePos="0" relativeHeight="0" 
                                   behindDoc="1" locked="0" layoutInCell="0" allowOverlap="1">
                            <wp:simplePos x="0" y="0"/>
                            <wp:positionH relativeFrom="page">
                                <wp:align>center</wp:align>
                            </wp:positionH>
                            <wp:positionV relativeFrom="page">
                                <wp:align>center</wp:align>
                            </wp:positionV>
                            <wp:extent cx="{cx}" cy="{cy}"/>
                            <wp:effectExtent l="0" t="0" r="0" b="0"/>
                            <wp:wrapNone/>
                        </wp:anchor>
                    '''
                    
                    anchor = parse_xml(anchor_xml)
                    
                    # Add docPr and graphic to anchor
                    anchor.append(docPr)
                    anchor.append(inline.cNvGraphicFramePr)
                    anchor.append(graphic)
                    
                    # Replace inline with anchor
                    drawing = picture._element
                    drawing_parent = drawing.getparent()
                    drawing_parent.replace(drawing, anchor)
                    
                except Exception as e2:
                    pass
        
        # Configure default font for the document to Calibri
        font_name = 'Calibri'
        
        # Set default font for the document
        style = doc.styles['Normal']
        style.font.name = font_name
        
        # Configure other default styles to use Calibri
        for style_name in ['Heading1', 'Heading2', 'Heading3', 'Heading4', 'Heading5', 'Heading6']:
            if style_name in doc.styles:
                doc.styles[style_name].font.name = font_name

        # Process HTML content
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "hr"]):
            if tag.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(tag.name[1])
                heading = doc.add_heading(tag.get_text().strip(), level=level)
                
                # Ensure heading uses Calibri
                for run in heading.runs:
                    run.font.name = font_name

            elif tag.name == "p":
                if tag.get_text().strip() == "":
                    continue

                paragraph = doc.add_paragraph()
                
                # Apply paragraph style attributes
                style = tag.get("style", "")

                if "text-align: center" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif "text-align: right" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif "text-align: left" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif "text-align: justify" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                if "padding-left" in style:
                    try:
                        padding_value = int(style.split("padding-left:")[1].split("px")[0].strip())
                        paragraph.paragraph_format.left_indent = Pt(padding_value)
                    except (ValueError, IndexError) as e:
                        pass

                if "line-height" in style:
                    try:
                        line_height = float(style.split("line-height:")[1].split(";")[0].strip())
                        paragraph.paragraph_format.line_spacing = line_height
                    except (ValueError, IndexError) as e:
                        pass

                # Define a helper function to apply styles to runs
                def apply_styles_to_run(run, element):
                    """Apply appropriate styles to a run based on the element and its style attributes"""
                    try:
                        from docx.shared import RGBColor
                        
                        element_name = element.name if hasattr(element, 'name') else None
                        element_style = element.get("style", "") if hasattr(element, 'get') else ""
                        
                        # Always apply Calibri font
                        run.font.name = font_name
                        
                        # Apply styles based on element type
                        if element_name == "strong" or element_name == "b":
                            run.bold = True
                        
                        if element_name == "em" or element_name == "i":
                            run.italic = True
                        
                        if element_name == "s" or "text-decoration: line-through" in element_style:
                            run.font.strike = True
                        
                        if element_name == "u" or "text-decoration: underline" in element_style:
                            run.underline = True
                        
                        # Apply span-specific styles
                        if element_name == "span":
                            if "text-decoration: underline" in element_style:
                                run.underline = True
                                
                            if "text-decoration: line-through" in element_style:
                                run.font.strike = True
                                
                            if "font-size" in element_style:
                                try:
                                    font_size_part = element_style.split("font-size:")[1].split(";")[0].strip()
                                    if "pt" in font_size_part:
                                        font_size = int(font_size_part.split("pt")[0].strip())
                                        run.font.size = Pt(font_size)
                                except (ValueError, IndexError) as e:
                                    pass
                                    
                            if "color:" in element_style or "color :" in element_style:
                                COLOR_MAP = {
                                    "red": (255, 0, 0),
                                    "green": (0, 128, 0),
                                    "blue": (0, 0, 255),
                                    "black": (0, 0, 0),
                                    "white": (255, 255, 255),
                                    "yellow": (255, 255, 0),
                                    "purple": (128, 0, 128),
                                    "orange": (255, 165, 0),
                                    "gray": (128, 128, 128),
                                    "pink": (255, 192, 203),
                                    "brown": (165, 42, 42),
                                    "cyan": (0, 255, 255),
                                    "magenta": (255, 0, 255),
                                    "lime": (0, 255, 0),
                                    "navy": (0, 0, 128),
                                    "teal": (0, 128, 128),
                                    "olive": (128, 128, 0),
                                    "maroon": (128, 0, 0),
                                    "silver": (192, 192, 192),
                                    "gold": (255, 215, 0)
                                }

                                try:
                                    normalized_style = element_style.replace(" :", ":")
                                    
                                    # Search color
                                    if "color:" in normalized_style:
                                        color_part = normalized_style.split("color:")[1].split(";")[0].strip()
                                    else:
                                        return run  # Color not found
                                    
                                    # Handle RGB colors
                                    if color_part.startswith("rgb("):
                                        color_values = color_part.replace("rgb(", "").replace(")", "").split(",")
                                        r = int(color_values[0].strip())
                                        g = int(color_values[1].strip())
                                        b = int(color_values[2].strip())
                                        run.font.color.rgb = RGBColor(r, g, b)
                                    
                                    # Handle color with name (red, blue, etc.)
                                    elif color_part in COLOR_MAP:
                                        r, g, b = COLOR_MAP[color_part]
                                        run.font.color.rgb = RGBColor(r, g, b)
                                    
                                    # Handle hexadecimal colors (#FF0000, etc.)
                                    elif color_part.startswith("#"):
                                        hex_color = color_part.lstrip("#")
                                        r = int(hex_color[0:2], 16)
                                        g = int(hex_color[2:4], 16)
                                        b = int(hex_color[4:6], 16)
                                        run.font.color.rgb = RGBColor(r, g, b)
                                        
                                except (ValueError, IndexError) as e:
                                    pass
                        
                        return run
                    except Exception as e:
                        return run

                # Improved recursive function to flatten the HTML structure
                def process_element_flat(element, current_styles=None):
                    """
                    Process elements by flattening the structure and tracking styles
                    This approach creates separate runs for each text node but applies all parent styles
                    """
                    from bs4 import NavigableString
                    
                    if current_styles is None:
                        current_styles = []
                    
                    # Skip None elements
                    if element is None:
                        return
                        
                    # For text nodes, create a run with all accumulated styles
                    if isinstance(element, NavigableString) and str(element).strip():
                        # Skip empty strings
                        if not str(element).strip():
                            return
                            
                        text = str(element)
                        
                        # Create a new run for this text
                        run = paragraph.add_run(text)
                        
                        # Apply all parent styles to this run
                        for style_element in current_styles:
                            run = apply_styles_to_run(run, style_element)
                            
                        # Ensure Calibri is applied even if no styles were applied
                        if not current_styles:
                            run.font.name = font_name
                            
                        return
                    
                    # If it's a tag element, add it to current styles and process children
                    if hasattr(element, 'name') and element.name:
                        # Add this element to the current style context
                        new_styles = current_styles + [element]
                        
                        # Process all children with updated styles
                        for child in element.children:
                            process_element_flat(child, new_styles)
                
                # Process paragraph using the flat approach
                for child in tag.children:
                    process_element_flat(child)

            elif tag.name == "hr":
                hr_paragraph = doc.add_paragraph("_" * 71)
                # Ensure Calibri is applied to the horizontal rule
                for run in hr_paragraph.runs:
                    run.font.name = font_name
        
        # Save the document to a buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return FileResponse(
            docx_buffer, 
            as_attachment=True, 
            filename=f"{document.title}.docx", 
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Document not found.'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'detail': f'Error generating Word document: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_recent_documents(request):
    """
    Get the 10 most recently visited documents for the authenticated user.
    Only returns documents the user has permission to view.
    """
    recent_documents = RecentDocument.objects.filter(user=request.user).select_related('document').order_by('-last_visited')
    
    # Filter by visibility permissions
    filtered_recent = []
    for recent_doc in recent_documents:
        if recent_doc.document.can_view(request.user):
            filtered_recent.append(recent_doc)
        if len(filtered_recent) >= 10:  # Limit to 10 documents
            break
    
    serializer = RecentDocumentSerializer(filtered_recent, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
@require_document_visibility_by_id
def update_recent_document(request, document_id):
    """
    Track a document visit by creating or updating a RecentDocument entry.
    Only allows tracking if user has permission to view the document.
    """
    try:
        document = DynamicDocument.objects.get(pk=document_id)
        recent_doc, created = RecentDocument.objects.get_or_create(
            user=request.user,
            document=document,
            defaults={'last_visited': timezone.now()}
        )
        
        if not created:
            recent_doc.last_visited = timezone.now()
            recent_doc.save()
            
        serializer = RecentDocumentSerializer(recent_doc)
        return Response(serializer.data, status=status.HTTP_200_OK)
        
    except DynamicDocument.DoesNotExist:
        return Response(
            {'detail': 'Document not found.'}, 
            status=status.HTTP_404_NOT_FOUND
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@require_document_usability('usability')
def upload_letterhead_image(request, pk):
    """
    Upload a letterhead image for a specific document.
    
    Accepts a PNG image file and saves it to the document's letterhead_image field.
    The image should be in 8.5:14 ratio (612x792 pixels recommended) for best results.
    
    Parameters:
        request (HttpRequest): The request object containing the image file
        pk (int): The primary key of the document
    
    Returns:
        Response: Success message with image info or error details
    """
    try:
        document = DynamicDocument.objects.get(pk=pk)
        
        # Check if image file is provided
        if 'image' not in request.FILES:
            return Response(
                {'detail': 'No se proporcionó ninguna imagen. Use el campo "image" para subir el archivo.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        image_file = request.FILES['image']
        
        # Validate file extension
        if not image_file.name.lower().endswith('.png'):
            return Response(
                {'detail': 'Solo se permiten archivos PNG para el membrete.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validate file size (max 10MB)
        max_size = 10 * 1024 * 1024  # 10MB
        if image_file.size > max_size:
            return Response(
                {'detail': 'El archivo es demasiado grande. Tamaño máximo: 10MB.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validate image dimensions and format using PIL
        try:
            img = Image.open(image_file)
            width, height = img.size
            
            # Check if it's a valid image
            img.verify()
            
            # Reset file pointer after verification
            image_file.seek(0)
            
            # Recommended dimensions check (warning, not error)
            recommended_width, recommended_height = 612, 792
            aspect_ratio = width / height
            recommended_ratio = recommended_width / recommended_height
            
            warnings = []
            if abs(aspect_ratio - recommended_ratio) > 0.1:  # Allow 10% tolerance
                warnings.append(f"La proporción de aspecto recomendada es 8.5:14 ({recommended_width}x{recommended_height}). Su imagen es {width}x{height}.")
            
        except Exception as e:
            return Response(
                {'detail': f'Archivo de imagen inválido: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Delete previous letterhead image if exists
        if document.letterhead_image:
            try:
                if os.path.exists(document.letterhead_image.path):
                    os.remove(document.letterhead_image.path)
            except:
                pass  # Continue even if deletion fails
        
        # Save the new image
        document.letterhead_image = image_file
        document.save(update_fields=['letterhead_image'])
        
        # Prepare response data
        response_data = {
            'message': 'Imagen de membrete subida exitosamente.',
            'document_id': document.id,
            'image_info': {
                'filename': image_file.name,
                'size_bytes': image_file.size,
                'dimensions': f"{width}x{height}",
                'format': img.format if hasattr(img, 'format') else 'PNG'
            }
        }
        
        if warnings:
            response_data['warnings'] = warnings
        
        return Response(response_data, status=status.HTTP_201_CREATED)
        
    except DynamicDocument.DoesNotExist:
        return Response(
            {'detail': 'Documento no encontrado.'},
            status=status.HTTP_404_NOT_FOUND
        )
    except Exception as e:
        return Response(
            {'detail': f'Error al subir la imagen: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@require_document_visibility
def get_letterhead_image(request, pk):
    """
    Get the letterhead image for a specific document.
    
    Parameters:
        request (HttpRequest): The request object
        pk (int): The primary key of the document
    
    Returns:
        FileResponse: The letterhead image file
        Response: Error message if image not found
    """
    try:
        document = DynamicDocument.objects.get(pk=pk)
        
        # Check if document has letterhead image
        if not document.letterhead_image:
            return Response(
                {'detail': 'Este documento no tiene imagen de membrete.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Check if file exists
        if not os.path.exists(document.letterhead_image.path):
            return Response(
                {'detail': 'El archivo de imagen no se encuentra en el servidor.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Return the image file
        return FileResponse(
            open(document.letterhead_image.path, 'rb'),
            as_attachment=False,
            filename=os.path.basename(document.letterhead_image.name),
            content_type='image/png'
        )
        
    except DynamicDocument.DoesNotExist:
        return Response(
            {'detail': 'Documento no encontrado.'},
            status=status.HTTP_404_NOT_FOUND
        )
    except Exception as e:
        return Response(
            {'detail': f'Error al obtener la imagen: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
@require_document_usability('usability')
def delete_letterhead_image(request, pk):
    """
    Delete the letterhead image for a specific document.
    
    Parameters:
        request (HttpRequest): The request object
        pk (int): The primary key of the document
    
    Returns:
        Response: Success or error message
    """
    try:
        document = DynamicDocument.objects.get(pk=pk)
        
        # Check if document has letterhead image
        if not document.letterhead_image:
            return Response(
                {'detail': 'Este documento no tiene imagen de membrete para eliminar.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Delete the file from filesystem
        try:
            if os.path.exists(document.letterhead_image.path):
                os.remove(document.letterhead_image.path)
        except Exception as e:
            pass
        
        # Clear the field
        document.letterhead_image = None
        document.save(update_fields=['letterhead_image'])
        
        return Response(
            {'message': 'Imagen de membrete eliminada exitosamente.'},
            status=status.HTTP_200_OK
        )
        
    except DynamicDocument.DoesNotExist:
        return Response(
            {'detail': 'Documento no encontrado.'},
            status=status.HTTP_404_NOT_FOUND
        )
    except Exception as e:
        return Response(
            {'detail': f'Error al eliminar la imagen: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


# ==================== LETTERHEAD HELPER FUNCTIONS ====================

def get_letterhead_for_document(document, user):
    """
    Get the appropriate letterhead image for a document.
    
    Priority:
    1. Document-specific letterhead (if exists)
    2. User's global letterhead (if exists)
    3. None (no letterhead)
    
    Args:
        document: DynamicDocument instance
        user: User instance (document creator/owner)
    
    Returns:
        ImageField or None: The letterhead image to use
    """
    # First priority: document-specific letterhead
    if document.letterhead_image:
        return document.letterhead_image
    
    # Second priority: user's global letterhead
    if user.letterhead_image:
        return user.letterhead_image
    
    # No letterhead available
    return None


# ==================== USER GLOBAL LETTERHEAD ENDPOINTS ====================

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def upload_user_letterhead_image(request):
    """
    Upload a global letterhead image for the authenticated user.
    
    This letterhead will be used as a fallback for all documents created by the user
    when they don't have a specific letterhead configured.
    
    Accepts a PNG image file and saves it to the user's letterhead_image field.
    The image should be in 8.5:11 ratio (612x792 pixels recommended) for best results.
    
    Parameters:
        request (HttpRequest): The request object containing the image file
    
    Returns:
        Response: Success message with image info or error message
    """
    # Restrict basic users from uploading letterhead images
    if request.user.role == 'basic':
        return Response({
            'error': 'Los usuarios básicos no pueden gestionar membretes'
        }, status=status.HTTP_403_FORBIDDEN)
    try:
        user = request.user
        
        # Check if image file is provided
        if 'image' not in request.FILES:
            return Response(
                {'detail': 'Se requiere un archivo de imagen.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        image_file = request.FILES['image']
        
        # Validate file extension
        if not image_file.name.lower().endswith('.png'):
            return Response(
                {'detail': 'Solo se permiten archivos PNG para el membrete.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validate file size (max 10MB)
        max_size = 10 * 1024 * 1024  # 10MB
        if image_file.size > max_size:
            return Response(
                {'detail': f'El archivo es demasiado grande. Tamaño máximo permitido: {max_size // (1024*1024)}MB.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validate image dimensions and format using PIL
        try:
            img = Image.open(image_file)
            width, height = img.size
            
            # Check if it's a valid image
            img.verify()
            
            # Reset file pointer after verification
            image_file.seek(0)
            
            # Recommended dimensions check (warning, not error)
            recommended_width, recommended_height = 612, 792
            aspect_ratio = width / height
            recommended_ratio = recommended_width / recommended_height
            
            warnings = []
            if abs(aspect_ratio - recommended_ratio) > 0.1:  # Allow 10% tolerance
                warnings.append(f"La proporción de aspecto recomendada es 8.5:11 ({recommended_width}x{recommended_height}). Su imagen es {width}x{height}.")
            
        except Exception as e:
            return Response(
                {'detail': f'Archivo de imagen inválido: {str(e)}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Delete previous letterhead image if exists
        if user.letterhead_image:
            try:
                if os.path.exists(user.letterhead_image.path):
                    os.remove(user.letterhead_image.path)
            except:
                pass  # Continue even if deletion fails
        
        # Save the new image
        user.letterhead_image = image_file
        user.save(update_fields=['letterhead_image'])
        
        # Prepare response data
        response_data = {
            'message': 'Membrete global subido exitosamente.',
            'user_id': user.id,
            'image_info': {
                'filename': image_file.name,
                'size_bytes': image_file.size,
                'dimensions': f"{width}x{height}",
                'format': img.format if hasattr(img, 'format') else 'PNG'
            }
        }
        
        if warnings:
            response_data['warnings'] = warnings
        
        return Response(response_data, status=status.HTTP_201_CREATED)
        
    except Exception as e:
        return Response(
            {'detail': f'Error al subir la imagen: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_user_letterhead_image(request):
    """
    Get the global letterhead image for the authenticated user.
    
    Returns:
        FileResponse: The letterhead image file
        Response: Error message if image not found
    """
    try:
        user = request.user
        
        # Check if user has letterhead image
        if not user.letterhead_image:
            return Response(
                {'detail': 'No tienes una imagen de membrete global configurada.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Check if file exists
        if not os.path.exists(user.letterhead_image.path):
            return Response(
                {'detail': 'El archivo de imagen no se encuentra en el servidor.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Return the image file
        return FileResponse(
            open(user.letterhead_image.path, 'rb'),
            as_attachment=False,
            filename=os.path.basename(user.letterhead_image.name),
            content_type='image/png'
        )
        
    except Exception as e:
        return Response(
            {'detail': f'Error al obtener la imagen: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_user_letterhead_image(request):
    """
    Delete the global letterhead image for the authenticated user.
    
    Returns:
        Response: Success or error message
    """
    # Restrict basic users from deleting letterhead images
    if request.user.role == 'basic':
        return Response({
            'error': 'Los usuarios básicos no pueden gestionar membretes'
        }, status=status.HTTP_403_FORBIDDEN)
    try:
        user = request.user
        
        # Check if user has letterhead image
        if not user.letterhead_image:
            return Response(
                {'detail': 'No tienes una imagen de membrete global para eliminar.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Delete the file from filesystem
        try:
            if os.path.exists(user.letterhead_image.path):
                os.remove(user.letterhead_image.path)
        except Exception as e:
            pass
        
        # Clear the field
        user.letterhead_image = None
        user.save(update_fields=['letterhead_image'])
        
        return Response(
            {'message': 'Membrete global eliminado exitosamente.'},
            status=status.HTTP_200_OK
        )
        
    except Exception as e:
        return Response(
            {'detail': f'Error al eliminar la imagen: {str(e)}'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )