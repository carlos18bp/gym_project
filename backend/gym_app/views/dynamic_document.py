import io
import re
from django.http import FileResponse
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from xhtml2pdf import pisa
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.template.loader import get_template
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from gym_app.models.dynamic_document import DynamicDocument
from gym_app.serializers.dynamic_document import DynamicDocumentSerializer


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def create_dynamic_document(request):
    """
    Create a new dynamic document.
    """
    print("Request data received:", request.data)

    # If it's a creation from the client, assign `assigned_to`.
    if request.data.get('state') in ['Progress', 'Completed'] and not request.data.get('assigned_to'):
        request.data['assigned_to'] = request.user.id

    # Validar y guardar el documento
    serializer = DynamicDocumentSerializer(data=request.data, context={'request': request})
    if serializer.is_valid():
        print("Serializer is valid. Saving document...")
        instance = serializer.save()
        print("Document saved successfully:", instance)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    print("Serializer errors:", serializer.errors)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def list_dynamic_documents(request):
    """
    Get a list of all dynamic documents.
    """
    documents = DynamicDocument.objects.prefetch_related('variables').all()
    serializer = DynamicDocumentSerializer(documents, many=True)
    return Response(serializer.data, status=status.HTTP_200_OK)


@api_view(['PUT', 'PATCH'])
@permission_classes([IsAuthenticated])
def update_dynamic_document(request, pk):
    """
    Update an existing dynamic document.
    """
    print(request.data)
    try:
        # Obtener el documento y cargar sus variables relacionadas
        document = DynamicDocument.objects.prefetch_related('variables').get(pk=pk)
    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Dynamic document not found.'}, status=status.HTTP_404_NOT_FOUND)

    # Evitar modificar el campo `created_by`
    if 'created_by' in request.data:
        request.data.pop('created_by')

    # Validar y actualizar el documento
    serializer = DynamicDocumentSerializer(document, data=request.data, partial=True, context={'request': request})
    if serializer.is_valid():
        print("Serializer is valid. Updating document...")
        instance = serializer.save()
        print("Document updated successfully:", instance)
        return Response(serializer.data, status=status.HTTP_200_OK)

    print("Serializer errors:", serializer.errors)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def delete_dynamic_document(request, pk):
    """
    Delete a dynamic document.
    """
    try:
        document = DynamicDocument.objects.get(pk=pk)
    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Dynamic document not found.'}, status=status.HTTP_404_NOT_FOUND)

    document.delete()
    print(f"Document with ID {pk} deleted successfully.")
    return Response({'detail': 'Dynamic document deleted successfully.'}, status=status.HTTP_200_OK)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_dynamic_document_pdf(request, pk):
    """
    Generates and returns a PDF file for the given document using ReportLab.
    """
    try:
        document = DynamicDocument.objects.prefetch_related('variables').get(pk=pk)

        # Replace variables in the content
        processed_content = document.content
        for variable in document.variables.all():
            processed_content = processed_content.replace(f"{{{{{variable.name_en}}}}}", variable.value or "")

        # Convert HTML to XHTML for proper parsing
        soup = BeautifulSoup(processed_content, 'html.parser')
        
        # Generate PDF using ReportLab and xhtml2pdf
        pdf_buffer = io.BytesIO()
        
        # Define the CSS styles
        styles = """
        <style>
        @page {
            margin: 2cm;
        }
        body {
            font-family: Helvetica, Arial, sans-serif;
            font-size: 12pt;
        }
        p {
            margin-bottom: 10px;
        }
        em {
            font-style: italic !important;
        }
        </style>
        """
        
        # Create the HTML with proper styling
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{document.title}</title>
            {styles}
        </head>
        <body>
            {soup.prettify()}
        </body>
        </html>
        """
        
        # Convert HTML to PDF
        pisa.CreatePDF(html_content, dest=pdf_buffer)
        pdf_buffer.seek(0)

        return FileResponse(
            pdf_buffer, 
            as_attachment=True, 
            filename=f"{document.title}.pdf", 
            content_type='application/pdf'
        )
    except DynamicDocument.DoesNotExist:
        return Response({'detail': 'Document not found.'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({'detail': f'Error generating PDF: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_dynamic_document_word(request, pk):
    """
    Generates and returns a Word (.docx) file for the given document using python-docx.
    """
    try:
        document = DynamicDocument.objects.prefetch_related('variables').get(pk=pk)

        # Replace variables dynamically
        def replace_variables(text):
            for variable in document.variables.all():
                pattern = re.compile(rf"{{{{{variable.name_en}}}}}")
                text = pattern.sub(variable.value or "", text)
            return text

        processed_content = replace_variables(document.content)
        
        # Render HTML with template
        template = get_template("pdf_template.html")
        html_content = template.render({"content": processed_content})

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        # Create Word document
        doc = Document()
        
        print("Starting document processing...")

        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "hr"]):
            if tag.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                level = int(tag.name[1])
                doc.add_heading(tag.get_text().strip(), level=level)

            elif tag.name == "p":
                if tag.get_text().strip() == "":
                    continue

                paragraph = doc.add_paragraph()
                
                # Apply paragraph style attributes
                style = tag.get("style", "")

                if "text-align: center" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif "text-align: right" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif "text-align: left" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                elif "text-align: justify" in style:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                if "padding-left" in style:
                    try:
                        padding_value = int(style.split("padding-left:")[1].split("px")[0].strip())
                        paragraph.paragraph_format.left_indent = Pt(padding_value)
                    except (ValueError, IndexError) as e:
                        print(f"  Error parsing padding-left: {str(e)}")

                if "line-height" in style:
                    try:
                        line_height = float(style.split("line-height:")[1].split(";")[0].strip())
                        paragraph.paragraph_format.line_spacing = line_height
                    except (ValueError, IndexError) as e:
                        print(f"  Error parsing line-height: {str(e)}")

                # Define a helper function to apply styles to runs
                def apply_styles_to_run(run, element):
                    """Apply appropriate styles to a run based on the element and its style attributes"""
                    try:
                        element_name = element.name if hasattr(element, 'name') else None
                        element_style = element.get("style", "") if hasattr(element, 'get') else ""
                        
                        
                        # Apply styles based on element type
                        if element_name == "strong" or element_name == "b":
                            run.bold = True
                        
                        if element_name == "em" or element_name == "i":
                            run.italic = True
                        
                        if element_name == "s" or "text-decoration: line-through" in element_style:
                            run.font.strike = True
                        
                        if element_name == "u" or "text-decoration: underline" in element_style:
                            run.underline = True
                        
                        # Apply span-specific styles
                        if element_name == "span":
                            if "text-decoration: underline" in element_style:
                                run.underline = True
                                
                            if "text-decoration: line-through" in element_style:
                                run.font.strike = True
                                
                            if "font-size" in element_style:
                                try:
                                    font_size_part = element_style.split("font-size:")[1].split(";")[0].strip()
                                    if "pt" in font_size_part:
                                        font_size = int(font_size_part.split("pt")[0].strip())
                                        run.font.size = Pt(font_size)
                                except (ValueError, IndexError) as e:
                                    print(f"    Error parsing font-size: {str(e)}")
                                    
                            if "color" in element_style:
                                try:
                                    color_part = element_style.split("color:")[1].split(";")[0].strip()
                                    if color_part.startswith("rgb("):
                                        color_values = color_part.replace("rgb(", "").replace(")", "").split(",")
                                        r = int(color_values[0].strip())
                                        g = int(color_values[1].strip())
                                        b = int(color_values[2].strip())
                                        run.font.color.rgb = RGBColor(r, g, b)
                                except (ValueError, IndexError) as e:
                                    print(f"    Error applying color: {str(e)}")
                        
                        return run
                    except Exception as e:
                        return run

                # Improved recursive function to flatten the HTML structure
                def process_element_flat(element, current_styles=None):
                    """
                    Process elements by flattening the structure and tracking styles
                    This approach creates separate runs for each text node but applies all parent styles
                    """
                    if current_styles is None:
                        current_styles = []
                    
                    # Skip None elements
                    if element is None:
                        return
                        
                    # For text nodes, create a run with all accumulated styles
                    if isinstance(element, NavigableString) and str(element).strip():
                        # Skip empty strings
                        if not str(element).strip():
                            return
                            
                        text = str(element)
                        
                        # Create a new run for this text
                        run = paragraph.add_run(text)
                        
                        # Apply all parent styles to this run
                        for style_element in current_styles:
                            run = apply_styles_to_run(run, style_element)
                            
                        return
                    
                    # If it's a tag element, add it to current styles and process children
                    if hasattr(element, 'name') and element.name:
                        # Add this element to the current style context
                        new_styles = current_styles + [element]
                        
                        # Process all children with updated styles
                        for child in element.children:
                            process_element_flat(child, new_styles)
                
                # Process paragraph using the flat approach
                for child in tag.children:
                    process_element_flat(child)

            elif tag.name == "hr":
                doc.add_paragraph("_" * 100)

        print("Document processing completed")
        
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return FileResponse(docx_buffer, as_attachment=True, filename=f"{document.title}.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except DynamicDocument.DoesNotExist:
        print("Error: Document not found")
        return Response({'detail': 'Document not found.'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        print(f"Error generating Word document: {str(e)}")
        return Response({'detail': f'Error generating Word document: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
